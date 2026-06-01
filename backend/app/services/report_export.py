from __future__ import annotations

import hashlib
from collections import defaultdict
from datetime import datetime
from io import BytesIO
from secrets import token_hex
from typing import Any
from zoneinfo import ZoneInfo, ZoneInfoNotFoundError

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.chart import BarChart, Reference
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.table import Table, TableStyleInfo

from app.core.config import get_settings
from app.services.dashboard import latest_report


DOCX_FONT_ASCII = "Noto Serif"
DOCX_FONT_EAST_ASIA = "Noto Serif CJK SC"
CHART_FONT_FAMILY = "Noto Serif CJK JP"

DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
XLSX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
ACCENT = "1F4E79"
ACCENT_DARK = "17365D"
ACCENT_LIGHT = "D9EAF7"
ACCENT_SOFT = "F2F6FB"
BORDER = "C7D6E8"
SUCCESS = "EAF4EC"
WARNING = "FFF2CC"
DANGER = "FCE4D6"
VM_ALERT_FILL = "F4CCCC"
VM_ALERT_RATIO = 0.2
VM_ALERT_BYTES = 100 * 1024**3
PERIODS = [("month", "较上月", 30), ("quarter", "较上季度", 90), ("year", "较上一年", 365)]


async def build_report_docx(tower_id: int | None = None, cluster_id: str | None = None, period_days: int = 30) -> tuple[bytes, str]:
    report = await latest_report(tower_id=tower_id, cluster_id=cluster_id, period_days=period_days)
    context = _export_context(report, tower_id, cluster_id)
    clusters = context["clusters"]
    document = Document()
    _setup_document(document, context)

    _add_cover(document, context)
    _add_cluster_directory(document, clusters)
    _add_report_overview(document, context)
    document.add_heading("1. 集群容量增长概览", level=1)
    _add_cluster_growth_chart(document, clusters)
    _add_cluster_summary_table(document, clusters)

    vms_by_cluster = _vms_by_cluster(report.get("month_fastest_growing_vms") or [])
    for index, cluster in enumerate(clusters, start=1):
        labels = cluster.get("labels", {})
        key = _cluster_key(labels)
        vms = vms_by_cluster.get(key, [])
        _start_section(document, context, _cluster_footer_label(labels))
        _add_bookmarked_heading(document, f"2.{index} {_cluster_title(labels)}", _cluster_bookmark(index), level=1, bookmark_id=index)
        _add_cluster_customer_summary(document, cluster, len(vms))
        _add_single_cluster_charts(document, cluster, vms)
        document.add_heading("增长量 TOP100 虚拟机（按增长量降序）", level=2)
        _add_vm_table(document, _top_vms(vms, "amount"), "amount")
        document.add_heading("增长率 TOP100 虚拟机（按增长率降序）", level=2)
        _add_vm_table(document, _top_vms(vms, "ratio"), "ratio")

    return _docx_bytes(document), _export_filename(context, "docx")


async def build_report_xlsx(tower_id: int | None = None, cluster_id: str | None = None, period_days: int = 30) -> tuple[bytes, str]:
    report = await latest_report(tower_id=tower_id, cluster_id=cluster_id, period_days=period_days)
    context = _export_context(report, tower_id, cluster_id)
    workbook = Workbook()
    summary_sheet = workbook.active
    summary_sheet.title = "汇总"
    _write_cluster_summary_sheet(summary_sheet, context)

    vms = report.get("month_fastest_growing_vms") or []
    _write_vm_summary_sheet(workbook.create_sheet("VM_TOP100_汇总"), vms)

    vms_by_cluster = _vms_by_cluster(vms)
    cluster_sheets: list[tuple[dict[str, Any], str]] = []
    for cluster in context["clusters"]:
        labels = cluster.get("labels", {})
        sheet = workbook.create_sheet(_safe_sheet_name(labels.get("cluster") or labels.get("cluster_id") or "集群"))
        cluster_sheets.append((cluster, sheet.title))
        cluster_vms = vms_by_cluster.get(_cluster_key(labels), [])
        _write_cluster_vm_sheet(sheet, cluster, cluster_vms)

    _write_directory_sheet(workbook.create_sheet("目录", 1), cluster_sheets)

    output = BytesIO()
    workbook.save(output)
    return output.getvalue(), _export_filename(context, "xlsx")


def _export_context(report: dict[str, Any], tower_id: int | None, cluster_id: str | None) -> dict[str, Any]:
    now = _local_now()
    clusters = report.get("clusters") or []
    if cluster_id:
        cluster_name = _first_label(clusters, "cluster") or cluster_id
        scope_slug = _slug(cluster_name)
        scope_label = f"集群 {cluster_name}"
    elif tower_id is not None:
        tower_name = _first_label(clusters, "tower") or str(tower_id)
        scope_slug = f"{_slug(tower_name)}-{token_hex(2)}"
        scope_label = f"Tower {tower_name}"
    else:
        scope_slug = "all"
        scope_label = "全部集群"
    return {
        "clusters": clusters,
        "date_slug": now.strftime("%Y%m%d"),
        "generated_at": now.strftime("%Y-%m-%d %H:%M:%S"),
        "timezone": _report_timezone_name(),
        "scope_label": scope_label,
        "scope_slug": scope_slug,
        "report": report,
    }


def _report_timezone_name() -> str:
    return get_settings().collection_timezone or "Asia/Shanghai"


def _local_now() -> datetime:
    timezone_name = _report_timezone_name()
    try:
        return datetime.now(ZoneInfo(timezone_name))
    except ZoneInfoNotFoundError:
        return datetime.now(ZoneInfo("Asia/Shanghai"))


def _export_filename(context: dict[str, Any], extension: str) -> str:
    return f"storage-forecast-{context['scope_slug']}-{context['date_slug']}.{extension}"


def _setup_document(document: Document, context: dict[str, Any]) -> None:
    section = document.sections[0]
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    section.top_margin = Inches(0.66)
    section.bottom_margin = Inches(0.64)
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = document.styles[style_name]
        _apply_style_font(style)
    document.styles["Normal"].font.size = Pt(9)
    document.styles["Heading 1"].font.size = Pt(15)
    document.styles["Heading 1"].font.bold = True
    document.styles["Heading 1"].font.color.rgb = RGBColor(23, 54, 93)
    document.styles["Heading 2"].font.size = Pt(11)
    document.styles["Heading 2"].font.bold = True
    document.styles["Heading 2"].font.color.rgb = RGBColor(31, 78, 121)
    _setup_header_footer(section, context)


def _apply_style_font(style: Any) -> None:
    style.font.name = DOCX_FONT_ASCII
    style._element.rPr.rFonts.set(qn("w:ascii"), DOCX_FONT_ASCII)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), DOCX_FONT_ASCII)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), DOCX_FONT_EAST_ASIA)


def _apply_run_font(run: Any) -> None:
    run.font.name = DOCX_FONT_ASCII
    run._element.rPr.rFonts.set(qn("w:ascii"), DOCX_FONT_ASCII)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), DOCX_FONT_ASCII)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), DOCX_FONT_EAST_ASIA)


def _setup_header_footer(section: Any, context: dict[str, Any], footer_label: str | None = None) -> None:
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("SmartX Storage Forecast | 容量预测报告")
    _apply_run_font(run)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(122, 138, 160)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f"{footer_label or context['scope_label']} · {context['generated_at']}")
    _apply_run_font(run)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(122, 138, 160)


def _start_section(document: Document, context: dict[str, Any], footer_label: str) -> None:
    section = document.add_section(WD_SECTION.NEW_PAGE)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    section.top_margin = Inches(0.66)
    section.bottom_margin = Inches(0.64)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    _clear_paragraph(section.header.paragraphs[0])
    _clear_paragraph(section.footer.paragraphs[0])
    _setup_header_footer(section, context, footer_label=footer_label)


def _clear_paragraph(paragraph: Any) -> None:
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)


def _add_cover_brand(document: Document) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_width(table, [900, 7200])
    _set_table_borders(table, "FFFFFF")
    _set_cell(table.rows[0].cells[0], "", fill=ACCENT)
    cell = table.rows[0].cells[1]
    _set_cell(cell, "SmartX HCI Capacity Insight", fill="F7FAFC", color=ACCENT_DARK, bold=True)
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_cover_rule(document: Document) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_width(table, [7800])
    _set_table_borders(table, "FFFFFF")
    _set_cell(table.rows[0].cells[0], "", fill=ACCENT)


def _add_callout(document: Document, title: str, body: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    _set_table_width(table, [7900])
    _set_table_borders(table, BORDER)
    cell = table.rows[0].cells[0]
    _shade_cell(cell, "F7FAFC")
    paragraph = cell.paragraphs[0]
    title_run = paragraph.add_run(f"{title}：")
    _apply_run_font(title_run)
    title_run.bold = True
    title_run.font.size = Pt(9)
    title_run.font.color.rgb = RGBColor(23, 54, 93)
    body_run = paragraph.add_run(body)
    _apply_run_font(body_run)
    body_run.font.size = Pt(9)
    body_run.font.color.rgb = RGBColor(38, 54, 79)
    document.add_paragraph()


def _add_figure(document: Document, image: BytesIO, caption: str, width: float) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(caption)
    _apply_run_font(run)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(61, 78, 99)
    paragraph.paragraph_format.space_after = Pt(2)
    picture = document.add_picture(image, width=Inches(width))
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()


def _add_cover(document: Document, context: dict[str, Any]) -> None:
    _add_cover_brand(document)
    document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("存储容量预测分析报告")
    _apply_run_font(run)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(23, 54, 93)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Storage Capacity Forecast Report")
    _apply_run_font(run)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(91, 115, 139)

    _add_cover_rule(document)
    document.add_paragraph()

    table = document.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    _set_table_width(table, [2300, 5400])
    _set_table_borders(table, BORDER)
    rows = [
        ("导出范围", context["scope_label"]),
        ("生成时间", f"{context['generated_at']} {context['timezone']}"),
        ("预测窗口", f"最近 {context['report'].get('window_days') or 30} 天，预测 {context['report'].get('forecast_days') or 60} 天"),
        ("集群数量", f"{len(context['clusters'])} 个"),
    ]
    for row, (key, value) in zip(table.rows, rows):
        _set_cell(row.cells[0], key, fill=ACCENT, color="FFFFFF", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell(row.cells[1], value, fill="F7FAFC")

    document.add_paragraph()
    _add_callout(
        document,
        "报告说明",
        "本报告基于平台采集容量样本自动生成，用于客户容量趋势复盘、风险沟通与扩容规划参考。预测结果应结合业务上线计划和实际资源治理策略共同判断。",
    )
    document.add_page_break()


def _add_report_overview(document: Document, context: dict[str, Any]) -> None:
    clusters = context["clusters"]
    total_current = sum(float((cluster.get("forecast") or {}).get("current") or 0) for cluster in clusters)
    total_month = sum(_cluster_period_growth(cluster, "month") for cluster in clusters)
    total_quarter = sum(_cluster_period_growth(cluster, "quarter") for cluster in clusters)
    total_year = sum(_cluster_period_growth(cluster, "year") for cluster in clusters)
    risk_count = sum(1 for cluster in clusters if _risk_level(cluster)[0] != "正常")
    document.add_heading("报告摘要", level=1)
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_width(table, [1650, 1650, 1650, 1650, 1650])
    _set_table_borders(table, BORDER)
    values = [
        ("当前已用", _format_bytes(total_current)),
        ("较上月增加", _format_bytes(total_month)),
        ("较上季度增加", _format_bytes(total_quarter)),
        ("较上一年增加", _format_bytes(total_year)),
        ("容量关注集群", f"{risk_count} 个"),
    ]
    for cell, (label, value) in zip(table.rows[0].cells, values):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _shade_cell(cell, ACCENT_SOFT)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(label + "\n")
        _apply_run_font(run)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(91, 115, 139)
        run = paragraph.add_run(value)
        _apply_run_font(run)
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(23, 54, 93)
    _add_callout(document, "结论摘要", _overview_sentence(clusters))


def _add_cluster_directory(document: Document, clusters: list[dict[str, Any]]) -> None:
    document.add_heading("目录", level=1)
    if not clusters:
        document.add_paragraph("当前导出范围内暂无集群。")
        document.add_page_break()
        return

    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_width(table, [900, 4300, 2500])
    _set_docx_headers(table.rows[0].cells, ["序号", "集群", "章节"])
    _set_table_borders(table, BORDER)
    for index, cluster in enumerate(clusters, start=1):
        labels = cluster.get("labels", {})
        row = table.add_row().cells
        _set_docx_row(row, [str(index), "", f"2.{index}"])
        _add_internal_link(row[1].paragraphs[0], _cluster_bookmark(index), _cluster_title(labels))
        if index % 2 == 0:
            _shade_row(row, "FAFCFF")
    _add_callout(document, "定位说明", "可通过本目录快速确认集群章节编号；Word 打开后也可使用导航窗格按标题跳转。")
    document.add_page_break()


def _overview_sentence(clusters: list[dict[str, Any]]) -> str:
    if not clusters:
        return "当前范围内暂无可用于分析的集群容量数据。"
    fastest = max(clusters, key=lambda item: _cluster_period_growth(item, "month"))
    risk_items = [cluster for cluster in clusters if _risk_level(cluster)[0] != "正常"]
    labels = fastest.get("labels", {})
    text = f"当前范围共 {len(clusters)} 个集群，月增长最高的是 {_cluster_title(labels)}，较上月增加 {_format_bytes(_cluster_period_growth(fastest, 'month'))}。"
    if risk_items:
        text += f" 有 {len(risk_items)} 个集群在 60 天预测窗口内达到或接近容量阈值，建议优先复核容量规划。"
    else:
        text += " 当前 60 天预测窗口内未发现明显容量阈值风险。"
    return text


def _add_cluster_growth_chart(document: Document, clusters: list[dict[str, Any]]) -> None:
    trend_chart = _scope_trend_line_chart(clusters, "容量使用率趋势")
    if trend_chart is not None:
        _add_figure(document, trend_chart, "图 1：容量使用趋势", width=6.6)
    else:
        document.add_paragraph("暂无足够历史数据生成容量趋势图。")

    top_chart = _cluster_top_growth_bar_chart(clusters, "Top 5 集群月增长量", period="month")
    if top_chart is not None:
        _add_figure(document, top_chart, "图 2：Top 5 集群月增长量", width=6.8)


def _add_cluster_summary_table(document: Document, clusters: list[dict[str, Any]]) -> None:
    headers = ["Tower", "集群", "当前容量", "较上月", "较上季度", "较上一年", "60天预测", "风险"]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_width(table, [1050, 1700, 1200, 1200, 1200, 1200, 1200, 850])
    _set_docx_headers(table.rows[0].cells, headers)
    _set_table_borders(table, BORDER)
    for cluster in clusters:
        labels = cluster.get("labels", {})
        forecast = cluster.get("forecast", {})
        risk, fill = _risk_level(cluster)
        row = table.add_row().cells
        values = [
            labels.get("tower") or labels.get("tower_id") or "-",
            labels.get("cluster") or labels.get("cluster_id") or "-",
            _format_bytes(forecast.get("current")),
            _format_bytes(_cluster_period_growth(cluster, "month")),
            _format_bytes(_cluster_period_growth(cluster, "quarter")),
            _format_bytes(_cluster_period_growth(cluster, "year")),
            _format_bytes(forecast.get("forecast_60d")),
            risk,
        ]
        _set_docx_row(row, values)
        if len(table.rows) % 2 == 1:
            _shade_row(row, "FAFCFF")
        _shade_cell(row[-1], fill)


def _add_cluster_customer_summary(document: Document, cluster: dict[str, Any], vm_count: int) -> None:
    forecast = cluster.get("forecast", {})
    risk, fill = _risk_level(cluster)
    table = document.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    _set_table_width(table, [2100, 2100, 2100, 2100])
    items = [
        ("当前已用", _format_bytes(forecast.get("current"))),
        ("较上月增加", _format_bytes(_cluster_period_growth(cluster, "month"))),
        ("较上季度增加", _format_bytes(_cluster_period_growth(cluster, "quarter"))),
        ("较上一年增加", _format_bytes(_cluster_period_growth(cluster, "year"))),
        ("60天预测", _format_bytes(forecast.get("forecast_60d"))),
        ("容量阈值", _format_bytes(cluster.get("warning"))),
        ("风险状态", risk),
        ("增长VM样本", f"{vm_count} 台"),
    ]
    for cell, (label, value) in zip([cell for row in table.rows for cell in row.cells], items):
        _shade_cell(cell, fill if label == "风险状态" else "F7FAFC")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(label + "\n")
        _apply_run_font(run)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(91, 115, 139)
        run = paragraph.add_run(value)
        _apply_run_font(run)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(23, 54, 93)
    document.add_paragraph(_cluster_advice(cluster))


def _cluster_advice(cluster: dict[str, Any]) -> str:
    risk, _ = _risk_level(cluster)
    month = _cluster_period_growth(cluster, "month")
    quarter = _cluster_period_growth(cluster, "quarter")
    if risk == "高风险":
        return "建议：当前容量已达到关注阈值，建议尽快确认可用容量、扩容周期和重点增长虚拟机。"
    if risk == "需关注":
        return "建议：60 天预测接近容量阈值，建议跟踪增长趋势并提前准备扩容或清理方案。"
    if quarter > month * 2.5 and quarter > 0:
        return "建议：季度增长明显高于月增长节奏，建议复核近期业务上线或批量数据写入情况。"
    return "建议：当前容量趋势相对平稳，建议持续观察 Top 增长虚拟机并保持月度复盘。"


def _add_single_cluster_charts(document: Document, cluster: dict[str, Any], vms: list[dict[str, Any]]) -> None:
    trend_chart = _cluster_trend_line_chart(cluster, "集群容量使用趋势")
    if trend_chart is not None:
        _add_figure(document, trend_chart, "容量使用趋势", width=5.9)

    top_chart = _vm_top_growth_bar_chart(vms, "Top 5 VM 增长量", mode="amount")
    if top_chart is not None:
        _add_figure(document, top_chart, "Top 5 VM 增长量", width=6.4)


def _add_vm_table(document: Document, vms: list[dict[str, Any]], sort_mode: str) -> None:
    headers = _vm_headers(include_cluster=False, sort_mode=sort_mode)
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _set_table_width(table, [650, 2850, 1200, 1200, 1200, 850])
    _set_docx_headers(table.rows[0].cells, headers)
    _set_table_borders(table, BORDER)
    if not vms:
        _set_docx_row(table.add_row().cells, ["-", "暂无增长数据", "-", "-", "-", "-"])
        return
    for index, vm in enumerate(vms, start=1):
        row = table.add_row().cells
        _set_docx_row(row, _vm_row(vm, include_cluster=False, rank=index))
        if _is_vm_alert(vm):
            _shade_row(row, VM_ALERT_FILL)
        elif index % 2 == 0:
            _shade_row(row, "FAFCFF")


def _set_docx_headers(cells: Any, headers: list[str]) -> None:
    for cell, header in zip(cells, headers):
        _set_cell(cell, header, fill=ACCENT, color="FFFFFF", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)


def _set_docx_row(cells: Any, values: list[str]) -> None:
    for cell, value in zip(cells, values):
        _set_cell(cell, value)


def _set_table_borders(table: Any, color: str) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _shade_row(cells: Any, fill: str) -> None:
    for cell in cells:
        _shade_cell(cell, fill)


def _set_cell(cell: Any, value: Any, fill: str | None = None, color: str = "1F1F1F", bold: bool = False, align: int | None = None) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        _shade_cell(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(str(value))
    _apply_run_font(run)
    run.bold = bold
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(color)
    for margin in ["top", "start", "bottom", "end"]:
        _set_cell_margin(cell, margin, 90)


def _add_bookmarked_heading(document: Document, text: str, bookmark_name: str, level: int, bookmark_id: int) -> None:
    paragraph = document.add_heading(level=level)
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), bookmark_name)
    paragraph._p.append(start)
    run = paragraph.add_run(text)
    _apply_run_font(run)
    if level == 1:
        run.bold = True
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(23, 54, 93)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.append(end)


def _add_internal_link(paragraph: Any, bookmark_name: str, text: str) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)
    run_element = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), ACCENT)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(color)
    properties.append(underline)
    run_element.append(properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run_element.append(text_element)
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)
    for run in paragraph.runs:
        _apply_run_font(run)


def _shade_cell(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margin(cell: Any, margin: str, size: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    node = tc_mar.find(qn(f"w:{margin}"))
    if node is None:
        node = OxmlElement(f"w:{margin}")
        tc_mar.append(node)
    node.set(qn("w:w"), str(size))
    node.set(qn("w:type"), "dxa")


def _set_table_width(table: Any, widths: list[int]) -> None:
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Pt(width / 20)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def _write_cluster_summary_sheet(sheet: Any, context: dict[str, Any]) -> None:
    sheet.append(["存储容量预测分析报告"])
    sheet.append(["导出范围", context["scope_label"]])
    sheet.append(["生成时间", f"{context['generated_at']} {context['timezone']}"])
    sheet.append(["预测窗口", f"最近 {context['report'].get('window_days') or 30} 天，预测 {context['report'].get('forecast_days') or 60} 天"])
    sheet.append([])
    headers = ["Tower", "集群", "当前容量", "较上月增加", "较上季度增加", "较上一年增加", "60天预测", "容量阈值", "耗尽天数", "风险"]
    sheet.append(headers)
    for cluster in context["clusters"]:
        labels = cluster.get("labels", {})
        forecast = cluster.get("forecast", {})
        risk, _ = _risk_level(cluster)
        sheet.append(
            [
                labels.get("tower") or labels.get("tower_id") or "-",
                labels.get("cluster") or labels.get("cluster_id") or "-",
                forecast.get("current") or 0,
                _cluster_period_growth(cluster, "month"),
                _cluster_period_growth(cluster, "quarter"),
                _cluster_period_growth(cluster, "year"),
                forecast.get("forecast_60d") or 0,
                cluster.get("warning") or 0,
                forecast.get("exhaustion_days") or "",
                risk,
            ]
        )
    _style_sheet(sheet, header_rows={1, 6}, bytes_cols={3, 4, 5, 6, 7, 8})
    _add_xlsx_growth_chart(sheet, start_row=6, end_row=sheet.max_row)


def _write_directory_sheet(sheet: Any, clusters: list[tuple[dict[str, Any], str]]) -> None:
    sheet.append(["集群目录"])
    sheet.append(["序号", "集群", "Sheet"])
    for index, (cluster, sheet_name) in enumerate(clusters, start=1):
        labels = cluster.get("labels", {})
        sheet.append([index, _cluster_title(labels), sheet_name])
        sheet.cell(row=sheet.max_row, column=3).hyperlink = f"#'{sheet_name}'!A1"
        sheet.cell(row=sheet.max_row, column=3).style = "Hyperlink"
    _style_sheet(sheet, header_rows={1, 2})


def _write_vm_summary_sheet(sheet: Any, vms: list[dict[str, Any]]) -> None:
    headers = _vm_headers(include_cluster=True, sort_mode="amount")
    _append_section_title(sheet, "增长量 TOP100（按增长量降序）", len(headers))
    amount_header_row = sheet.max_row + 1
    sheet.append(headers)
    amount_start = sheet.max_row + 1
    for index, vm in enumerate(_top_vms(vms, "amount"), start=1):
        sheet.append(_vm_row(vm, include_cluster=True, raw=True, rank=index))
    amount_end = sheet.max_row

    sheet.append([])
    ratio_title_row = sheet.max_row + 1
    _append_section_title(sheet, "增长率 TOP100（按增长率降序）", len(headers))
    ratio_header_row = sheet.max_row + 1
    sheet.append(_vm_headers(include_cluster=True, sort_mode="ratio"))
    ratio_start = sheet.max_row + 1
    for index, vm in enumerate(_top_vms(vms, "ratio"), start=1):
        sheet.append(_vm_row(vm, include_cluster=True, raw=True, rank=index))
    ratio_end = sheet.max_row

    header_rows = {1, amount_header_row, ratio_title_row, ratio_header_row}
    _style_sheet(sheet, header_rows=header_rows, bytes_cols={5, 6, 7}, percent_cols={8})
    _style_vm_rows(sheet, amount_start, amount_end, include_cluster=True, column_count=len(headers))
    _style_vm_rows(sheet, ratio_start, ratio_end, include_cluster=True, column_count=len(headers))
    _add_excel_table(sheet, "VmAmountSummary", amount_header_row, amount_end, len(headers))
    _add_excel_table(sheet, "VmRatioSummary", ratio_header_row, ratio_end, len(headers))


def _write_cluster_vm_sheet(sheet: Any, cluster: dict[str, Any], vms: list[dict[str, Any]]) -> None:
    labels = cluster.get("labels", {})
    forecast = cluster.get("forecast", {})
    risk, _ = _risk_level(cluster)
    sheet.append([_cluster_title(labels)])
    sheet.append(["当前容量", forecast.get("current") or 0, "较上月增加", _cluster_period_growth(cluster, "month"), "较上季度增加", _cluster_period_growth(cluster, "quarter"), "较上一年增加", _cluster_period_growth(cluster, "year"), "风险", risk])
    sheet.append([])
    sheet.append(["增长量 TOP100（按增长量降序）"])
    amount_headers = _vm_headers(include_cluster=False, sort_mode="amount")
    sheet.append(amount_headers)
    amount_header_row = sheet.max_row
    amount_start = sheet.max_row + 1
    for index, vm in enumerate(_top_vms(vms, "amount"), start=1):
        sheet.append(_vm_row(vm, include_cluster=False, raw=True, rank=index))
    amount_end = sheet.max_row
    start = sheet.max_row + 2
    sheet.cell(row=start, column=1, value="增长率 TOP100（按增长率降序）")
    ratio_headers = _vm_headers(include_cluster=False, sort_mode="ratio")
    for col, header in enumerate(ratio_headers, start=1):
        sheet.cell(row=start + 1, column=col, value=header)
    ratio_start = start + 2
    for rank, (row_index, vm) in enumerate(zip(range(start + 2, start + 2 + len(_top_vms(vms, "ratio"))), _top_vms(vms, "ratio")), start=1):
        for col, value in enumerate(_vm_row(vm, include_cluster=False, raw=True, rank=rank), start=1):
            sheet.cell(row=row_index, column=col, value=value)
    ratio_end = sheet.max_row
    _style_sheet(sheet, header_rows={1, 2, 4, 5, start, start + 1}, bytes_cols={3, 4, 5, 7, 9}, percent_cols={6})
    _style_vm_rows(sheet, amount_start, amount_end, include_cluster=False, column_count=len(amount_headers))
    _style_vm_rows(sheet, ratio_start, ratio_end, include_cluster=False, column_count=len(ratio_headers))
    _add_excel_table(sheet, _table_safe_name(sheet.title, "Amount"), amount_header_row, amount_end, len(amount_headers))
    _add_excel_table(sheet, _table_safe_name(sheet.title, "Ratio"), start + 1, ratio_end, len(ratio_headers))


def _style_sheet(sheet: Any, header_rows: set[int], bytes_cols: set[int] | None = None, percent_cols: set[int] | None = None) -> None:
    bytes_cols = bytes_cols or set()
    percent_cols = percent_cols or set()
    fill = PatternFill(fill_type="solid", fgColor=ACCENT_LIGHT)
    title_fill = PatternFill(fill_type="solid", fgColor=ACCENT)
    for row in sheet.iter_rows():
        for cell in row:
            cell.alignment = Alignment(vertical="center", wrap_text=True)
            if cell.row == 1:
                cell.font = Font(bold=True, color="FFFFFF", size=13)
                cell.fill = title_fill
            elif cell.row in header_rows:
                cell.font = Font(bold=True, color=ACCENT_DARK)
                cell.fill = fill
            if cell.column in bytes_cols and isinstance(cell.value, (int, float)):
                cell.number_format = '#,##0'
            if cell.column in percent_cols and isinstance(cell.value, (int, float)):
                cell.number_format = "0.00%"
    for column in range(1, sheet.max_column + 1):
        max_len = 10
        for cell in sheet[get_column_letter(column)]:
            max_len = max(max_len, len(str(cell.value or "")))
        sheet.column_dimensions[get_column_letter(column)].width = min(max_len + 2, 32)
    sheet.freeze_panes = "A2"


def _append_section_title(sheet: Any, title: str, columns: int) -> None:
    sheet.append([title])
    if columns > 1:
        sheet.merge_cells(start_row=sheet.max_row, start_column=1, end_row=sheet.max_row, end_column=columns)


def _style_vm_rows(sheet: Any, start_row: int, end_row: int, include_cluster: bool, column_count: int) -> None:
    if end_row < start_row:
        return
    fill = PatternFill(fill_type="solid", fgColor=VM_ALERT_FILL)
    for row_index in range(start_row, end_row + 1):
        amount_col = 7 if include_cluster else 5
        ratio_col = 8 if include_cluster else 6
        amount = sheet.cell(row=row_index, column=amount_col).value
        ratio = sheet.cell(row=row_index, column=ratio_col).value
        if _is_vm_alert_values(amount, ratio):
            for column in range(1, column_count + 1):
                sheet.cell(row=row_index, column=column).fill = fill


def _add_excel_table(sheet: Any, name: str, header_row: int, end_row: int, column_count: int) -> None:
    if end_row <= header_row:
        return
    ref = f"A{header_row}:{get_column_letter(column_count)}{end_row}"
    table = Table(displayName=name[:255], ref=ref)
    table.tableStyleInfo = TableStyleInfo(name="TableStyleMedium2", showRowStripes=True, showColumnStripes=False)
    sheet.add_table(table)


def _add_xlsx_growth_chart(sheet: Any, start_row: int, end_row: int) -> None:
    if end_row <= start_row:
        return
    chart = BarChart()
    chart.type = "bar"
    chart.style = 10
    chart.title = "集群容量增长对比"
    chart.y_axis.title = "容量增长"
    chart.x_axis.title = "集群"
    data = Reference(sheet, min_col=4, max_col=6, min_row=start_row, max_row=end_row)
    cats = Reference(sheet, min_col=2, min_row=start_row + 1, max_row=end_row)
    chart.add_data(data, titles_from_data=True)
    chart.set_categories(cats)
    chart.height = 8
    chart.width = 18
    sheet.add_chart(chart, "L2")


def _scope_trend_line_chart(clusters: list[dict[str, Any]], title: str) -> BytesIO | None:
    points = _merged_cluster_points(clusters)
    return _line_chart_image(points, title=title, ylabel="容量 (TiB)")


def _cluster_trend_line_chart(cluster: dict[str, Any], title: str) -> BytesIO | None:
    return _line_chart_image(_cluster_points(cluster), title=title, ylabel="容量 (TiB)")


def _line_chart_image(points: list[tuple[int, float]], title: str, ylabel: str) -> BytesIO | None:
    if len(points) < 2:
        return None
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.dates as mdates
        import matplotlib.pyplot as plt
        matplotlib.rcParams["font.family"] = [CHART_FONT_FAMILY, "Noto Serif", "DejaVu Serif"]
        matplotlib.rcParams["axes.unicode_minus"] = False
    except Exception:
        return None
    dates = [datetime.fromtimestamp(ts) for ts, _ in points]
    values = [_bytes_to_tib(value) for _, value in points]
    fig, ax = plt.subplots(figsize=(6.6, 4.0), dpi=180)
    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")
    ax.plot(dates, values, color="#003BFF", linewidth=1.4)
    ax.set_title(title, fontsize=12, pad=8)
    ax.set_xlabel("时间", fontsize=10)
    ax.set_ylabel(ylabel, fontsize=10)
    y_min, y_max = _chart_y_limits(values)
    ax.set_ylim(y_min, y_max)
    ax.grid(axis="y", color="#E6EDF5", linewidth=0.6)
    ax.xaxis.set_major_locator(mdates.AutoDateLocator(minticks=4, maxticks=6))
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%Y-%m-%d"))
    for label in ax.get_xticklabels():
        label.set_rotation(0)
        label.set_ha("center")
    for spine in ax.spines.values():
        spine.set_color("#9E9E9E")
        spine.set_linewidth(0.9)
    ax.tick_params(axis="both", labelsize=9, colors="#333333")
    fig.tight_layout()
    return _figure_bytes(fig, plt)


def _chart_y_limits(values: list[float]) -> tuple[float, float]:
    if not values:
        return 0.0, 1.0
    y_min = min(values)
    y_max = max(values)
    if y_min == y_max:
        padding = max(abs(y_max) * 0.05, 0.1)
    else:
        padding = max((y_max - y_min) * 0.18, y_max * 0.01, 0.1)
    lower = max(0.0, y_min - padding)
    upper = y_max + padding
    if upper <= lower:
        upper = lower + 1.0
    return lower, upper


def _cluster_top_growth_bar_chart(clusters: list[dict[str, Any]], title: str, period: str) -> BytesIO | None:
    items = []
    for cluster in clusters:
        labels = cluster.get("labels", {})
        name = labels.get("cluster") or labels.get("cluster_id") or "集群"
        value = _cluster_period_growth(cluster, period)
        if value > 0:
            items.append((str(name), value))
    return _horizontal_bar_chart_image(items, title=title, unit="TiB", limit=5)


def _vm_top_growth_bar_chart(vms: list[dict[str, Any]], title: str, mode: str) -> BytesIO | None:
    ranked = _top_vms(vms, "ratio" if mode == "ratio" else "amount")[:5]
    items = []
    for vm in ranked:
        labels = vm.get("labels", {})
        name = labels.get("vm") or labels.get("vm_id") or "VM"
        value = (vm.get("growth_ratio") or 0) if mode == "ratio" else (vm.get("growth_amount") or 0)
        if value:
            items.append((str(name), float(value)))
    return _horizontal_bar_chart_image(items, title=title, unit="%" if mode == "ratio" else "TiB", limit=5, percent=mode == "ratio")


def _horizontal_bar_chart_image(items: list[tuple[str, float]], title: str, unit: str, limit: int = 5, percent: bool = False) -> BytesIO | None:
    if not items:
        return None
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        matplotlib.rcParams["font.family"] = [CHART_FONT_FAMILY, "Noto Serif", "DejaVu Serif"]
        matplotlib.rcParams["axes.unicode_minus"] = False
    except Exception:
        return None
    ranked = sorted(items, key=lambda item: item[1], reverse=True)[:limit]
    names = [_truncate_label(name) for name, _ in ranked][::-1]
    raw_values = [value for _, value in ranked][::-1]
    values = [value * 100 if percent else _bytes_to_tib(value) for value in raw_values]
    max_value = max(values) if values else 0
    fig_height = max(2.2, 0.45 * len(values) + 1.1)
    fig, ax = plt.subplots(figsize=(6.8, fig_height), dpi=180)
    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")
    bars = ax.barh(names, values, color="#1155CC", height=0.32)
    ax.set_title(title, fontsize=12, pad=8)
    ax.set_xlim(0, max(max_value * 1.18, 1))
    ax.grid(axis="x", color="#E6EDF5", linewidth=0.6)
    for spine in ax.spines.values():
        spine.set_color("#9E9E9E")
        spine.set_linewidth(0.9)
    ax.tick_params(axis="both", labelsize=9, colors="#333333")
    for bar, value in zip(bars, values):
        label = f"{value:.1f} %" if percent else f"{value:.2f} {unit}"
        ax.text(bar.get_width(), bar.get_y() + bar.get_height() / 2, label, va="center", ha="left", fontsize=9, color="#333333")
    fig.tight_layout()
    return _figure_bytes(fig, plt)


def _figure_bytes(fig: Any, plt: Any) -> BytesIO:
    output = BytesIO()
    fig.savefig(output, format="png", bbox_inches="tight", facecolor="white")
    plt.close(fig)
    output.seek(0)
    return output


def _merged_cluster_points(clusters: list[dict[str, Any]]) -> list[tuple[int, float]]:
    by_ts: dict[int, float] = defaultdict(float)
    for cluster in clusters:
        for ts, value in _cluster_points(cluster):
            by_ts[ts] += value
    return sorted(by_ts.items())


def _truncate_label(value: str, limit: int = 22) -> str:
    text = str(value)
    if len(text) <= limit:
        return text
    return f"{text[:8]}...{text[-8:]}"


def _cluster_period_growth(cluster: dict[str, Any], period: str) -> float:
    days_by_period = {key: days for key, _, days in PERIODS}
    days = days_by_period.get(period, 30)
    points = _cluster_points(cluster)
    forecast = cluster.get("forecast", {})
    current = float(forecast.get("current") or (points[-1][1] if points else 0) or 0)
    if len(points) >= 2:
        return max(0.0, current - _value_days_ago(points, days))
    return max(0.0, float(forecast.get("slope_per_day") or 0) * days)


def _cluster_points(cluster: dict[str, Any]) -> list[tuple[int, float]]:
    points = []
    for point in cluster.get("points") or []:
        try:
            points.append((int(float(point[0])), float(point[1])))
        except (TypeError, ValueError, IndexError):
            continue
    return sorted(points)


def _value_days_ago(points: list[tuple[int, float]], days: int) -> float:
    if not points:
        return 0.0
    target = points[-1][0] - days * 86_400
    candidates = [point for point in points if point[0] <= target]
    if candidates:
        return candidates[-1][1]
    return points[0][1]


def _risk_level(cluster: dict[str, Any]) -> tuple[str, str]:
    forecast = cluster.get("forecast", {})
    current = float(forecast.get("current") or 0)
    future = float(forecast.get("forecast_60d") or current)
    warning = float(cluster.get("warning") or 0)
    if warning and current >= warning:
        return "高风险", DANGER
    if warning and future >= warning:
        return "需关注", WARNING
    return "正常", SUCCESS


def _vms_by_cluster(vms: list[dict[str, Any]]) -> dict[tuple[str, str], list[dict[str, Any]]]:
    grouped: dict[tuple[str, str], list[dict[str, Any]]] = defaultdict(list)
    for vm in vms:
        grouped[_cluster_key(vm.get("labels", {}))].append(vm)
    return grouped


def _top_vms(vms: list[dict[str, Any]], mode: str) -> list[dict[str, Any]]:
    key = (lambda vm: vm.get("growth_ratio") or 0) if mode == "ratio" else (lambda vm: vm.get("growth_amount") or 0)
    return sorted(vms, key=key, reverse=True)[:100]


def _vm_headers(include_cluster: bool, sort_mode: str) -> list[str]:
    headers = ["排名", "VM", "当前容量", "上期容量", "增长量", "增长率"]
    if sort_mode == "amount":
        headers[4] = "增长量 ↓"
    elif sort_mode == "ratio":
        headers[5] = "增长率 ↓"
    if include_cluster:
        return ["Tower", "集群", *headers]
    return headers


def _vm_row(vm: dict[str, Any], include_cluster: bool, raw: bool = False, rank: int | None = None) -> list[Any]:
    labels = vm.get("labels", {})
    forecast = vm.get("forecast", {})
    ratio = vm.get("growth_ratio")
    values: list[Any] = []
    if include_cluster:
        values.extend([labels.get("tower") or labels.get("tower_id") or "-", labels.get("cluster") or labels.get("cluster_id") or "-"])
    values.extend(
        [
            rank or "",
            labels.get("vm") or labels.get("vm_id") or "-",
            forecast.get("current") or 0,
            vm.get("previous_value") or 0,
            vm.get("growth_amount") or 0,
            ratio if ratio is not None else "",
        ]
    )
    if raw:
        return values
    formatted = [str(values[0]), str(values[1])]
    formatted.extend(_format_bytes(value) for value in values[2:5])
    formatted.append(_format_percent(ratio))
    return formatted


def _is_vm_alert(vm: dict[str, Any]) -> bool:
    return _is_vm_alert_values(vm.get("growth_amount"), vm.get("growth_ratio"))


def _is_vm_alert_values(amount: Any, ratio: Any) -> bool:
    try:
        amount_value = float(amount or 0)
    except (TypeError, ValueError):
        amount_value = 0.0
    try:
        ratio_value = float(ratio or 0)
    except (TypeError, ValueError):
        ratio_value = 0.0
    return amount_value > VM_ALERT_BYTES and ratio_value > VM_ALERT_RATIO


def _cluster_title(labels: dict[str, Any]) -> str:
    tower = labels.get("tower") or labels.get("tower_id") or "Tower"
    cluster = labels.get("cluster") or labels.get("cluster_id") or "集群"
    return f"{tower} / {cluster}"


def _cluster_footer_label(labels: dict[str, Any]) -> str:
    tower = str(labels.get("tower") or labels.get("tower_id") or "Tower")
    cluster = str(labels.get("cluster") or labels.get("cluster_id") or "集群")
    return f"{tower}-{cluster}集群"


def _cluster_key(labels: dict[str, Any]) -> tuple[str, str]:
    return (str(labels.get("tower_id") or ""), str(labels.get("cluster_id") or ""))


def _cluster_bookmark(index: int) -> str:
    return f"cluster_{index}"


def _monthly_growth(cluster: dict[str, Any]) -> float:
    return _cluster_period_growth(cluster, "month")


def _bytes_to_tib(value: Any) -> float:
    try:
        return float(value or 0) / 1024**4
    except (TypeError, ValueError):
        return 0.0


def _format_bytes(value: Any) -> str:
    try:
        size = float(value or 0)
    except (TypeError, ValueError):
        size = 0
    if size <= 0:
        return "0 B"
    units = ["B", "KiB", "MiB", "GiB", "TiB", "PiB"]
    index = 0
    while size >= 1024 and index < len(units) - 1:
        size /= 1024
        index += 1
    return f"{size:.0f} {units[index]}" if index < 3 else f"{size:.2f} {units[index]}"


def _format_percent(value: Any) -> str:
    try:
        if value is None:
            return "-"
        numeric = float(value)
    except (TypeError, ValueError):
        return "-"
    return f"{numeric * 100:.2f}%"


def _format_days(value: Any) -> str:
    try:
        if value is None:
            return "未触发"
        return f"{float(value):.0f} 天"
    except (TypeError, ValueError):
        return "未触发"


def _first_label(clusters: list[dict[str, Any]], key: str) -> str | None:
    for cluster in clusters:
        value = cluster.get("labels", {}).get(key)
        if value:
            return str(value)
    return None


def _safe_sheet_name(value: str) -> str:
    invalid = set('[]:*?/\\')
    cleaned = "".join("_" if char in invalid else char for char in str(value))
    return (cleaned or "集群")[:31]


def _table_safe_name(value: str, suffix: str = "") -> str:
    digest = hashlib.sha1(f"{value}:{suffix}".encode("utf-8")).hexdigest()[:8]
    return f"T_{digest}{suffix}"[:60]


def _slug(value: str) -> str:
    return "".join(char if char.isalnum() or char in "-_" else "-" for char in value)[:48]


def _docx_bytes(document: Document) -> bytes:
    output = BytesIO()
    document.save(output)
    return output.getvalue()
