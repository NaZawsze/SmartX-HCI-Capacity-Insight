from __future__ import annotations

from collections import defaultdict
from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
from pathlib import Path
import re
from typing import Any
from urllib.parse import quote
from zoneinfo import ZoneInfo, ZoneInfoNotFoundError

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl.chart import BarChart, Reference
from openpyxl import Workbook, load_workbook
from openpyxl.cell.cell import MergedCell
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.table import Table, TableStyleInfo

from app.v2.config import V2Settings


DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
XLSX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
VM_ALERT_FILL = "F4CCCC"
VM_ALERT_RATIO = 0.2
VM_ALERT_BYTES = 100 * 1024**3
DOCX_FONT_ASCII = "Noto Serif"
DOCX_FONT_EAST_ASIA = "Noto Serif CJK SC"
CHART_FONT_FAMILY = "Noto Serif CJK JP"
CHART_FONT_PATH = "/usr/share/fonts/opentype/noto/NotoSerifCJK-Regular.ttc"
REPORT_PRODUCT_NAME = "存储容量预测平台"
REPORT_COVER_TITLE = "SMARTX超融合存储容量分析报告"
REPORT_COVER_SUBTITLE = "SMARTX HCI Storage Capacity Analysis Report"
XLSX_TEMPLATE_PATH = Path(__file__).with_name("templates") / "customer_report.xlsx"
ACCENT = "1A3C6E"
ACCENT_DARK = "1A3C6E"
ACCENT_LIGHT = "F0F4FA"
ACCENT_SOFT = "F7F9FC"
BORDER = "D9E2EF"
SUCCESS = "EAF7EF"
WARNING = "FFF4E5"
DANGER = "FDEAEA"
TEXT_DARK = "333333"
TEXT_MUTED = "999999"
GROWTH_BLUE = "007ACC"
RATIO_ORANGE = "E67E22"
RATIO_RED = "DC3535"
EMPTY_MONTH_VM_TEXT = "当前样本跨度区间暂无 VM 增长数据"
PERIODS = [("month", "较上月", 30), ("quarter", "较上季度", 90), ("year", "较上一年", 365)]


@dataclass(frozen=True)
class ReportPeriodProfile:
    days: int
    display_label: str
    window_growth_label: str
    vm_growth_title: str
    vm_empty_text: str
    short_advice_title: str
    focus_note: str


REPORT_PERIOD_PROFILES: dict[int, ReportPeriodProfile] = {
    7: ReportPeriodProfile(
        days=7,
        display_label="近 7 天",
        window_growth_label="近 7 天样本增长",
        vm_growth_title="短期突增 VM",
        vm_empty_text="当前 7 天窗口暂无明显 VM 增长数据",
        short_advice_title="短期（本周）",
        focus_note="本报告侧重识别短期突增、本日新建 VM 和日增长来源。",
    ),
    14: ReportPeriodProfile(
        days=14,
        display_label="近 14 天",
        window_growth_label="近 14 天样本增长",
        vm_growth_title="近两周增长 VM",
        vm_empty_text="当前 14 天窗口暂无明显 VM 增长数据",
        short_advice_title="短期（两周内）",
        focus_note="本报告侧重观察近两周增长变化和短期异常来源。",
    ),
    30: ReportPeriodProfile(
        days=30,
        display_label="近 30 天",
        window_growth_label="近 30 天样本增长",
        vm_growth_title="月度增长 VM",
        vm_empty_text="当前 30 天窗口暂无明显 VM 增长数据",
        short_advice_title="短期（本月）",
        focus_note="本报告侧重月度运营窗口、容量增长和 VM TOP 变化。",
    ),
    90: ReportPeriodProfile(
        days=90,
        display_label="近 90 天",
        window_growth_label="近 90 天样本增长",
        vm_growth_title="季度增长 VM",
        vm_empty_text="当前 90 天窗口暂无明显 VM 增长数据",
        short_advice_title="短期（本季度）",
        focus_note="本报告侧重季度趋势、预测可信度和容量风险。",
    ),
    180: ReportPeriodProfile(
        days=180,
        display_label="近 180 天",
        window_growth_label="近 180 天样本增长",
        vm_growth_title="中长期增长 VM",
        vm_empty_text="当前 180 天窗口暂无明显 VM 增长数据",
        short_advice_title="短期（半年内）",
        focus_note="本报告侧重中长期增长趋势和容量治理。",
    ),
    365: ReportPeriodProfile(
        days=365,
        display_label="近 365 天",
        window_growth_label="近 365 天样本增长",
        vm_growth_title="年度增长 VM",
        vm_empty_text="当前 365 天窗口暂无明显 VM 增长数据",
        short_advice_title="短期（年度巡检）",
        focus_note="本报告侧重年度容量规划；采集历史不足一年时按可用样本窗口计算。",
    ),
}


def report_period_profile(period_days: int | None) -> ReportPeriodProfile:
    try:
        days = int(period_days or 30)
    except (TypeError, ValueError):
        days = 30
    return REPORT_PERIOD_PROFILES.get(days, REPORT_PERIOD_PROFILES[30])


def build_report_docx(report: dict[str, Any], settings: V2Settings, *, period_days: int) -> tuple[bytes, str, Path, str]:
    context = _export_context(report, settings, period_days, "docx")
    clusters = report.get("clusters") or []
    month_vms = report.get("month_fastest_growing_vms") or []
    document = Document()
    _setup_document(document)
    _setup_footer(document, report, context)
    _add_cover(document, report, context, settings)
    _add_cluster_directory(document, clusters)
    document.add_heading("报告摘要", level=1)
    _add_paragraph_table(
        document,
        [
            ("导出范围", context["scope_label"]),
            ("生成时间", context["generated_at"]),
            ("统计窗口", _period_window_label(report)),
            ("预测窗口", f"{report.get('forecast_days', 90)} 天"),
            ("集群数量", str(len(report.get("clusters") or []))),
            ("容量风险摘要", _capacity_risk_summary(report)),
            ("当前软件版本", settings.app_version),
        ],
    )
    _add_callout(document, "结论摘要", _overview_sentence(clusters))
    document.add_heading("集群容量增长概览", level=1)
    _add_cluster_growth_chart(document, clusters)
    _add_cluster_table(document, clusters)
    document.add_heading("日增长最快 VM", level=1)
    _add_vm_table(document, report.get("day_fastest_growing_vms") or [], empty_text="暂无日增长 VM 数据", include_cluster=True)
    document.add_heading("本日新建 VM", level=1)
    _add_new_vm_table(document, report.get("day_new_vms") or [], empty_text="暂无本日新建 VM")
    document.add_heading("本月新建 VM", level=1)
    _add_new_vm_table(document, report.get("month_new_vms") or [], empty_text="暂无本月新建 VM")
    document.add_heading("VM_TOP100_汇总", level=1)
    document.add_heading("增长量 TOP100", level=2)
    document.add_paragraph(f"统计窗口：{_period_window_label(report)}")
    _add_vm_table(document, _top_vms(month_vms, "amount"), empty_text=EMPTY_MONTH_VM_TEXT, include_cluster=True)
    document.add_heading("增长率 TOP100", level=2)
    document.add_paragraph(f"统计窗口：{_period_window_label(report)}")
    _add_vm_table(document, _top_vms(month_vms, "ratio"), empty_text=EMPTY_MONTH_VM_TEXT, include_cluster=True)
    document.add_heading("集群专项分析", level=1)
    figure_index = 3
    for index, cluster in enumerate(clusters, start=1):
        labels = cluster.get("labels", {})
        cluster_vms = _vms_by_cluster(month_vms).get(_cluster_key(labels), [])
        document.add_page_break()
        document.add_heading(f"{index}. {_cluster_name(cluster)} 集群容量概览", level=1)
        _add_single_cluster_summary(document, cluster, len(cluster_vms))
        figure_index = _add_single_cluster_charts(document, cluster, cluster_vms, figure_index=figure_index)
        document.add_heading("增长量 TOP100 虚拟机（按增长量降序）", level=2)
        document.add_paragraph(f"统计窗口：{_period_window_label(report)}")
        _add_vm_table(document, _top_vms(cluster_vms, "amount"), empty_text=EMPTY_MONTH_VM_TEXT, include_cluster=False)
        document.add_heading("增长率 TOP100 虚拟机（按增长率降序）", level=2)
        document.add_paragraph(f"统计窗口：{_period_window_label(report)}")
        _add_vm_table(document, _top_vms(cluster_vms, "ratio"), empty_text=EMPTY_MONTH_VM_TEXT, include_cluster=False)
    content = _docx_bytes(document)
    return _persist_report(content, settings, context["filename"])


def build_report_xlsx(report: dict[str, Any], settings: V2Settings, *, period_days: int) -> tuple[bytes, str, Path, str]:
    context = _export_context(report, settings, period_days, "xlsx")
    clusters = report.get("clusters") or []
    month_vms = report.get("month_fastest_growing_vms") or []
    profile = context["profile"]
    workbook = _load_customer_xlsx_template()
    _write_xlsx_template_cover(workbook["封面"], report, context, settings)
    _write_xlsx_template_summary(workbook["执行摘要"], report, context, clusters, settings, profile)
    _write_xlsx_template_capacity_trend(workbook["容量趋势"], report, context, clusters, profile)
    top_sheet = workbook["VM增长TOP20"] if "VM增长TOP20" in workbook.sheetnames else _replace_sheet(workbook, "VM增长TOP100")
    top_sheet.title = "VM增长TOP100"
    _write_xlsx_template_vm_top100(top_sheet, _customer_growth_vms(report), report, profile)
    _write_xlsx_template_day_growth(workbook["日增长详情"], report.get("day_fastest_growing_vms") or [], report)

    _write_simple_vm_sheet(_replace_sheet(workbook, "本日新建VM"), report.get("day_new_vms") or [], "暂无本日新建 VM", include_growth=False)
    _write_simple_vm_sheet(_replace_sheet(workbook, "本月新建VM"), report.get("month_new_vms") or [], "暂无本月新建 VM", include_growth=False)
    vms_by_cluster = _vms_by_cluster(month_vms)
    for cluster in clusters:
        labels = cluster.get("labels", {})
        sheet = _replace_sheet(workbook, _safe_sheet_name(labels.get("cluster") or labels.get("cluster_id") or "集群"))
        _write_cluster_vm_sheet(sheet, cluster, vms_by_cluster.get(_cluster_key(labels), []), report, profile)
    _remove_sheets(workbook, ["目录", "范围明细", "集群汇总", "VM_TOP100_汇总", "汇总"])

    output = BytesIO()
    workbook.save(output)
    return _persist_report(output.getvalue(), settings, context["filename"])


def _setup_document(document: Document) -> None:
    section = document.sections[0]
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    section.top_margin = Inches(0.66)
    section.bottom_margin = Inches(0.64)
    styles = document.styles
    styles["Normal"].font.name = DOCX_FONT_ASCII
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = document.styles[style_name]
        style.font.name = DOCX_FONT_ASCII
        style._element.rPr.rFonts.set(qn("w:ascii"), DOCX_FONT_ASCII)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), DOCX_FONT_ASCII)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), DOCX_FONT_EAST_ASIA)


def _add_cover(document: Document, report: dict[str, Any], context: dict[str, str], settings: V2Settings) -> None:
    document.add_heading(REPORT_COVER_TITLE, level=0)
    subtitle = document.add_paragraph(REPORT_COVER_SUBTITLE)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor.from_string(ACCENT_DARK)
    _add_paragraph_table(
        document,
        [
            ("导出范围", context["scope_label"]),
            ("生成时间", context["generated_at"]),
            ("统计窗口", _period_window_label(report)),
            ("预测窗口", f"{report.get('forecast_days', 90)} 天"),
            ("集群数量", str(len(report.get("clusters") or []))),
            ("当前软件版本", settings.app_version),
        ],
    )
    _add_callout(document, "报告说明", "本报告基于平台采集容量样本自动生成，用于客户容量趋势复盘、风险沟通与扩容规划参考。")
    document.add_page_break()


def _add_callout(document: Document, title: str, body: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    _shade_cell(cell, ACCENT_SOFT)
    paragraph = cell.paragraphs[0]
    paragraph.add_run(f"{title}：").bold = True
    paragraph.add_run(body)


def _add_paragraph_table(document: Document, rows: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "项目"
    table.rows[0].cells[1].text = "内容"
    _shade_cell(table.rows[0].cells[0], ACCENT)
    _shade_cell(table.rows[0].cells[1], ACCENT)
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value


def _add_cluster_directory(document: Document, clusters: list[dict[str, Any]]) -> None:
    document.add_heading("报告正文目录", level=1)
    if not clusters:
        document.add_paragraph("暂无集群目录")
        return
    document.add_paragraph("1. 报告摘要")
    document.add_paragraph("2. 集群容量增长概览")
    document.add_paragraph("3. VM 增长与新建明细")
    document.add_paragraph("4. 集群专项分析")
    for index, cluster in enumerate(clusters, start=1):
        labels = cluster.get("labels", {})
        name = labels.get("cluster") or labels.get("cluster_id") or f"集群 {index}"
        document.add_paragraph(f"4.{index} {name}", style="List Number")


def _add_cluster_table(document: Document, clusters: list[dict[str, Any]]) -> None:
    table = document.add_table(rows=1, cols=8)
    table.style = "Table Grid"
    headers = ["Tower", "集群", "当前容量", "较上月", "较上季度", "较上一年", "90 天预测", "风险"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
        _shade_cell(table.rows[0].cells[index], ACCENT)
    if not clusters:
        row = table.add_row().cells
        row[0].text = "当前范围暂无集群容量数据"
        return
    for cluster in clusters:
        labels = cluster.get("labels", {})
        forecast = cluster.get("forecast", {})
        risk, fill = _risk_level(cluster)
        values = [
            labels.get("tower") or labels.get("tower_id") or "",
            labels.get("cluster") or labels.get("cluster_id") or "",
            _bytes_label(forecast.get("current")),
            _bytes_label(_cluster_period_growth(cluster, 30)),
            _bytes_label(_cluster_period_growth(cluster, 90)),
            _bytes_label(_cluster_period_growth(cluster, 365)),
            _bytes_label(forecast.get("forecast_90d")),
            risk,
        ]
        row = table.add_row().cells
        for index, value in enumerate(values):
            row[index].text = value
        _shade_cell(row[-1], fill)


def _add_vm_table(document: Document, vms: list[dict[str, Any]], *, empty_text: str, include_cluster: bool = True) -> None:
    headers = ["Tower", "集群", "VM", "当前容量", "期初容量", "增长量", "增长率"] if include_cluster else ["VM", "当前容量", "期初容量", "增长量", "增长率"]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
        _shade_cell(table.rows[0].cells[index], ACCENT)
    if not vms:
        row = table.add_row().cells
        row[0].text = empty_text
        return
    for vm in vms[:100]:
        labels = vm.get("labels", {})
        forecast = vm.get("forecast", {})
        values = []
        if include_cluster:
            values.extend([labels.get("tower") or labels.get("tower_id") or "", labels.get("cluster") or labels.get("cluster_id") or ""])
        values.extend([
            labels.get("vm") or labels.get("vm_name") or labels.get("vm_id") or "",
            _bytes_label(forecast.get("current")),
            _bytes_label(vm.get("previous_value")),
            _bytes_label(vm.get("growth_amount")),
            _percent_label(vm.get("growth_ratio")),
        ])
        row = table.add_row().cells
        for index, value in enumerate(values):
            row[index].text = value
        if _is_alert_vm(vm):
            for cell in row:
                _shade_cell(cell, VM_ALERT_FILL)


def _add_new_vm_table(document: Document, vms: list[dict[str, Any]], *, empty_text: str) -> None:
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["Tower", "集群", "VM", "当前容量", "首次出现时间"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
        _shade_cell(table.rows[0].cells[index], ACCENT)
    if not vms:
        table.add_row().cells[0].text = empty_text
        return
    for vm in vms[:100]:
        labels = vm.get("labels", {})
        row = table.add_row().cells
        values = [
            labels.get("tower") or labels.get("tower_id") or "",
            labels.get("cluster") or labels.get("cluster_id") or "",
            labels.get("vm") or labels.get("vm_id") or "",
            _bytes_label((vm.get("forecast") or {}).get("current")),
            str(vm.get("first_seen_at") or ""),
        ]
        for index, value in enumerate(values):
            row[index].text = value


def _export_context(report: dict[str, Any], settings: V2Settings, period_days: int, extension: str) -> dict[str, str]:
    now = _local_now(settings)
    profile = report_period_profile(period_days)
    scope = report.get("scope") or {}
    clusters = report.get("clusters") or []
    if scope.get("cluster_id"):
        scope_label = _first_cluster_name(clusters) or str(scope["cluster_id"])
    elif scope.get("tower_id") is not None:
        scope_label = f"tower-{scope['tower_id']}"
    else:
        scope_label = "all"
    scope_slug = _slug(scope_label)
    filename = f"storage-forecast-{scope_slug}-{now.strftime('%Y%m%d-%H%M%S')}-{period_days}d.{extension}"
    return {
        "filename": filename,
        "scope_label": scope_label,
        "generated_at": now.strftime("%Y-%m-%d %H:%M:%S %Z"),
        "clusters": clusters,
        "report": report,
        "period_window": report.get("period_window") or {},
        "app_version": settings.app_version,
        "profile": profile,
    }


def _persist_report(content: bytes, settings: V2Settings, filename: str) -> tuple[bytes, str, Path, str]:
    settings.reports_dir.mkdir(parents=True, exist_ok=True)
    path = settings.reports_dir / Path(filename).name
    path.write_bytes(content)
    return content, filename, path, f"/api/admin/exports/reports/{quote(path.name)}"


def _setup_footer(document: Document, report: dict[str, Any], context: dict[str, str]) -> None:
    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.text = _footer_label(report, context)


def _footer_label(report: dict[str, Any], context: dict[str, str]) -> str:
    clusters = report.get("clusters") or []
    tower_names = sorted({str((cluster.get("labels") or {}).get("tower") or "") for cluster in clusters if (cluster.get("labels") or {}).get("tower")})
    cluster_names = [str((cluster.get("labels") or {}).get("cluster") or (cluster.get("labels") or {}).get("cluster_id") or "") for cluster in clusters]
    tower_label = "、".join(tower_names[:2]) if tower_names else context["scope_label"]
    cluster_label = "、".join([name for name in cluster_names if name][:3]) if cluster_names else "全部集群"
    if len(cluster_names) > 3:
        cluster_label += f" 等 {len(cluster_names)} 个集群"
    return f"{tower_label} - {cluster_label} - {context['generated_at']}"


def _overview_sentence(clusters: list[dict[str, Any]]) -> str:
    if not clusters:
        return "当前范围内暂无可用于分析的集群容量数据。"
    fastest = max(clusters, key=lambda item: _cluster_period_growth(item, 30))
    risk_items = [cluster for cluster in clusters if _risk_level(cluster)[0] != "正常"]
    text = f"当前范围共 {len(clusters)} 个集群，月增长最高的是 {_cluster_name(fastest)}，较上月增加 {_bytes_label(_cluster_period_growth(fastest, 30))}。"
    if risk_items:
        text += f" 有 {len(risk_items)} 个集群在 90 天预测窗口内达到或接近容量阈值，建议优先复核容量规划。"
    else:
        text += " 当前 90 天预测窗口内未发现明显容量阈值风险。"
    return text


def _add_cluster_growth_chart(document: Document, clusters: list[dict[str, Any]]) -> None:
    trend_chart = _scope_trend_line_chart(clusters, "容量使用率趋势")
    if trend_chart is not None:
        _add_figure(document, trend_chart, "图 1：容量使用趋势", width=6.4)
    else:
        document.add_paragraph("暂无足够历史数据生成容量趋势图。")
    top_chart = _cluster_top_growth_bar_chart(clusters, "Top 5 集群月增长量")
    if top_chart is not None:
        _add_figure(document, top_chart, "图 2：Top 5 集群月增长量", width=6.5)


def _add_single_cluster_summary(document: Document, cluster: dict[str, Any], vm_count: int) -> None:
    forecast = cluster.get("forecast", {})
    risk, fill = _risk_level(cluster)
    rows = [
        ("当前容量", _bytes_label(forecast.get("current"))),
        ("较上月增加", _bytes_label(_cluster_period_growth(cluster, 30))),
        ("较上季度增加", _bytes_label(_cluster_period_growth(cluster, 90))),
        ("较上一年增加", _bytes_label(_cluster_period_growth(cluster, 365))),
        ("90 天预测容量", _bytes_label(forecast.get("forecast_90d"))),
        ("容量阈值", _bytes_label(cluster.get("warning"))),
        ("风险状态", risk),
        ("增长 VM 样本", f"{vm_count} 台"),
    ]
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "项目"
    table.rows[0].cells[1].text = "内容"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
        if key == "风险状态":
            _shade_cell(cells[1], fill)


def _add_single_cluster_charts(document: Document, cluster: dict[str, Any], vms: list[dict[str, Any]], *, figure_index: int) -> int:
    cluster_name = _cluster_name(cluster)
    trend_chart = _cluster_trend_line_chart(cluster, "集群容量使用趋势")
    if trend_chart is not None:
        _add_figure(document, trend_chart, f"图 {figure_index}：{cluster_name} 容量使用趋势", width=5.9)
    else:
        document.add_paragraph(f"图 {figure_index}：{cluster_name} 容量使用趋势：暂无足够历史数据生成图表。")
    figure_index += 1
    top_chart = _vm_top_growth_bar_chart(vms, "Top 10 VM 增长量")
    if top_chart is not None:
        _add_figure(document, top_chart, f"图 {figure_index}：{cluster_name} Top 10 VM 增长量", width=6.4)
    else:
        document.add_paragraph(f"图 {figure_index}：{cluster_name} Top 10 VM 增长量：{EMPTY_MONTH_VM_TEXT}。")
    return figure_index + 1


def _add_figure(document: Document, image: BytesIO, caption: str, width: float) -> None:
    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption_paragraph.add_run(caption)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(ACCENT_DARK)
    picture_paragraph = document.add_paragraph()
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.add_run().add_picture(image, width=Inches(width))


def _period_window_label(report: dict[str, Any]) -> str:
    window = _effective_report_window(report)
    start = str(window.get("start_at") or "")
    end = str(window.get("end_at") or "")
    if start and end:
        timezone = _report_timezone(report)
        parsed_start = _parse_report_datetime(start)
        parsed_end = _parse_report_datetime(end)
        if parsed_start and parsed_end:
            return f"{parsed_start.astimezone(timezone).date().isoformat()} - {parsed_end.astimezone(timezone).date().isoformat()}"
        return f"{start[:10]} - {end[:10]}"
    return f"近 {report.get('window_days', 30)} 天"


def _requested_report_window_label(report: dict[str, Any]) -> str:
    period = report.get("period_window") or {}
    days = period.get("days") or report.get("window_days") or 30
    start = str(period.get("start_at") or "")
    end = str(period.get("end_at") or "")
    try:
        days_value = int(days)
    except (TypeError, ValueError):
        days_value = 30
    if start and end:
        timezone = _report_timezone(report)
        parsed_start = _parse_report_datetime(start)
        parsed_end = _parse_report_datetime(end)
        if parsed_start and parsed_end:
            return (
                f"近 {days_value} 天（"
                f"{parsed_start.astimezone(timezone).date().isoformat()} - "
                f"{parsed_end.astimezone(timezone).date().isoformat()}）"
            )
        return f"近 {days_value} 天（{start[:10]} - {end[:10]}）"
    return f"近 {days_value} 天"


def _effective_report_window(report: dict[str, Any]) -> dict[str, Any]:
    timezone = _report_timezone(report)
    period = report.get("period_window") or {}
    data = report.get("data_window") or {}
    period_start = _parse_report_datetime(str(period.get("start_at") or ""))
    period_end = _parse_report_datetime(str(period.get("end_at") or ""))
    data_start = _parse_report_datetime(str(data.get("start_at") or ""))
    data_end = _parse_report_datetime(str(data.get("end_at") or ""))
    start_candidates = [value for value in [period_start, data_start] if value is not None]
    end_candidates = [value for value in [period_end, data_end] if value is not None]
    if start_candidates and end_candidates:
        start = max(start_candidates)
        end = min(end_candidates)
        if start <= end:
            return {"start_at": start.astimezone(timezone).isoformat(), "end_at": end.astimezone(timezone).isoformat()}
    if data_start and data_end:
        return {"start_at": data_start.astimezone(timezone).isoformat(), "end_at": data_end.astimezone(timezone).isoformat()}
    if period_start and period_end:
        return {"start_at": period_start.astimezone(timezone).isoformat(), "end_at": period_end.astimezone(timezone).isoformat()}
    return {}


def _vm_sample_window_label(vms: list[dict[str, Any]], report: dict[str, Any]) -> str:
    starts = [_parse_report_datetime(str(vm.get("window_start_at") or "")) for vm in vms]
    ends = [_parse_report_datetime(str(vm.get("window_end_at") or "")) for vm in vms]
    starts = [value for value in starts if value is not None]
    ends = [value for value in ends if value is not None]
    if starts and ends:
        timezone = _report_timezone(report)
        return f"{min(starts).astimezone(timezone).date().isoformat()} - {max(ends).astimezone(timezone).date().isoformat()}"
    return _period_window_label(report)


def _parse_report_datetime(value: str) -> datetime | None:
    if not value:
        return None
    try:
        return datetime.fromisoformat(value.replace("Z", "+00:00"))
    except ValueError:
        return None


def _report_timezone(report: dict[str, Any]) -> ZoneInfo:
    timezone_name = str(report.get("timezone") or "Asia/Shanghai")
    try:
        return ZoneInfo(timezone_name)
    except ZoneInfoNotFoundError:
        return ZoneInfo("Asia/Shanghai")


def _capacity_risk_summary(report: dict[str, Any]) -> str:
    clusters = report.get("clusters") or []
    if not clusters:
        return "暂无集群容量数据，无法判断容量风险。"
    ranked = sorted(clusters, key=_cluster_used_ratio, reverse=True)
    high = [cluster for cluster in ranked if _cluster_used_ratio(cluster) >= 0.8]
    warning = [cluster for cluster in ranked if _cluster_used_ratio(cluster) >= 0.75]
    if high:
        return _risk_summary_sentence(high, "使用率超过 80%，容量风险较高")
    if warning:
        return _risk_summary_sentence(warning, "使用率超过 75%，需要关注容量增长")
    return "当前所有集群暂无明显容量风险。"


def _risk_summary_sentence(clusters: list[dict[str, Any]], suffix: str) -> str:
    names = "、".join(_cluster_name(cluster) for cluster in clusters[:3])
    if len(clusters) > 3:
        names += f" 等 {len(clusters)} 个集群"
    return f"{names} {suffix}。"


def _cluster_name(cluster: dict[str, Any]) -> str:
    labels = cluster.get("labels") or {}
    return str(labels.get("cluster") or labels.get("cluster_id") or "未知集群")


def _cluster_full_name(cluster: dict[str, Any]) -> str:
    labels = cluster.get("labels") or {}
    tower = labels.get("tower") or labels.get("tower_id")
    cluster_name = labels.get("cluster") or labels.get("cluster_id") or "未知集群"
    return f"{tower} / {cluster_name}" if tower else str(cluster_name)


def _tower_scope_label(clusters: list[dict[str, Any]], fallback: str) -> str:
    names = []
    for cluster in clusters:
        labels = cluster.get("labels") or {}
        value = labels.get("tower") or labels.get("tower_id")
        if value and str(value) not in names:
            names.append(str(value))
    if not names:
        return fallback
    if len(names) == 1:
        return names[0]
    return f"全部 Tower（{len(names)} 个）"


def _cluster_scope_label(clusters: list[dict[str, Any]], fallback: str) -> str:
    names = []
    for cluster in clusters:
        name = _cluster_full_name(cluster)
        if name and name not in names:
            names.append(name)
    if not names:
        return fallback
    if len(names) == 1:
        return names[0]
    return f"全部集群（{len(names)} 个）"


def _cluster_used_ratio(cluster: dict[str, Any]) -> float:
    total = _float_or_none(cluster.get("total"))
    current = _float_or_none((cluster.get("forecast") or {}).get("current"))
    if total and total > 0 and current is not None:
        return current / total
    return 0.0


def _risk_level(cluster: dict[str, Any]) -> tuple[str, str]:
    used_ratio = _cluster_used_ratio(cluster)
    forecast = cluster.get("forecast") or {}
    future = _float_or_none(forecast.get("forecast_90d")) or _float_or_none(forecast.get("current")) or 0
    warning = _float_or_none(cluster.get("warning")) or 0
    if used_ratio >= 0.8:
        return "高风险", DANGER
    if warning and future >= warning:
        return "需关注", WARNING
    return "正常", SUCCESS


def _overall_risk_status(clusters: list[dict[str, Any]]) -> tuple[str, str, str]:
    if not clusters:
        return "数据不足", WARNING, "当前导出范围暂无集群容量数据。"
    high = [cluster for cluster in clusters if _cluster_used_ratio(cluster) >= 0.8]
    warning = [cluster for cluster in clusters if _cluster_used_ratio(cluster) >= 0.75 or _exhaustion_days(cluster) <= 180]
    if high:
        return "高风险", DANGER, f"{_cluster_full_name(high[0])} 容量使用率已达到高风险阈值。"
    if warning:
        return "需关注", WARNING, f"{_cluster_full_name(warning[0])} 容量增长或预测耗尽时间需要关注。"
    return "正常", SUCCESS, "当前集群整体运行平稳，暂无明显容量风险。"


def _customer_growth_vms(report: dict[str, Any]) -> list[dict[str, Any]]:
    return _merge_growth_candidates(report.get("window_fastest_growing_vms") or [], report.get("month_fastest_growing_vms") or [])


def _report_vm_count(report: dict[str, Any]) -> int:
    return len(_report_vm_keys(report))


def _report_cluster_vm_counts(report: dict[str, Any]) -> dict[tuple[str, str], int]:
    grouped: dict[tuple[str, str], set[str]] = defaultdict(set)
    for tower_id, cluster_id, vm_id in _report_vm_keys(report):
        grouped[(tower_id, cluster_id)].add(vm_id)
    return {key: len(vm_ids) for key, vm_ids in grouped.items()}


def _cluster_vm_count(cluster: dict[str, Any], fallback: int | None) -> int:
    labels = cluster.get("labels") or {}
    value = cluster.get("vm_count") or cluster.get("virtual_machine_count") or cluster.get("vms")
    try:
        numeric = int(value)
        if numeric >= 0:
            return numeric
    except (TypeError, ValueError):
        pass
    return int(fallback or 0)


def _report_vm_keys(report: dict[str, Any]) -> set[tuple[str, str, str]]:
    keys: set[tuple[str, str, str]] = set()
    for key in ("day_fastest_growing_vms", "window_fastest_growing_vms", "month_fastest_growing_vms", "day_new_vms", "month_new_vms"):
        for vm in report.get(key) or []:
            labels = vm.get("labels") or {}
            vm_key = (
                str(labels.get("tower_id") or ""),
                str(labels.get("cluster_id") or ""),
                str(labels.get("vm_id") or labels.get("vm") or labels.get("vm_name") or ""),
            )
            if any(vm_key):
                keys.add(vm_key)
    return keys


def _customer_key_findings(clusters: list[dict[str, Any]], top_vms: list[dict[str, Any]], risk_note: str, context: dict[str, Any]) -> list[str]:
    if not clusters:
        return ["当前导出范围暂无可用于分析的集群容量数据，建议确认 Tower 与集群采集状态。"]
    total_current = sum(float((cluster.get("forecast") or {}).get("current") or 0) for cluster in clusters)
    total_capacity = sum(float(cluster.get("total") or 0) for cluster in clusters)
    total_window = sum(_cluster_period_growth(cluster, _customer_window_days(context["report"])) for cluster in clusters)
    total_forecast_90 = sum(float((cluster.get("forecast") or {}).get("forecast_90d") or 0) for cluster in clusters)
    profile = context["profile"]
    findings = [
        f"当前导出范围已用容量 {_bytes_label(total_current)}，{profile.window_growth_label} {_bytes_label(total_window)}，{risk_note}",
    ]
    if total_capacity > 0:
        findings.append(f"按当前增长趋势推算，90 天后容量预计为 {_bytes_label(total_forecast_90)}，约占总容量 {_bytes_label(total_capacity)} 的 {_percent_label(total_forecast_90 / total_capacity)}。")
    if top_vms:
        first = top_vms[0]
        findings.append(
            f"{profile.display_label}内，{_vm_scope_name(first)} 下的 {_vm_display_name(first)} "
            f"增长最为显著，增长量 {_bytes_label(first.get('growth_amount'))}，建议确认增长来源的合理性。"
        )
    largest_vm = _largest_vm(top_vms)
    if largest_vm:
        findings.append(
            f"{_vm_scope_name(largest_vm)} 下的 {_vm_display_name(largest_vm)} "
            f"当前容量 {_bytes_label((largest_vm.get('forecast') or {}).get('current'))}，建议纳入重点监控清单。"
        )
    high_ratio = [vm for vm in top_vms if float(vm.get("growth_ratio") or 0) >= 1.0]
    if high_ratio:
        findings.append(f"增长率超过 100% 的 VM 有 {len(high_ratio)} 台，可能属于基数较小导致的增长率偏高，需结合绝对值综合评估。")
    return findings[:5]


def _customer_risk_matrix_rows(clusters: list[dict[str, Any]], top_vms: list[dict[str, Any]]) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    if clusters:
        worst = max(clusters, key=_cluster_used_ratio)
        ratio = _cluster_used_ratio(worst)
        level = "高" if ratio >= 0.8 else "中" if ratio >= 0.75 else "低"
        rows.append({
            "item": "集群整体容量",
            "status": _risk_level(worst)[0],
            "level": level,
            "description": f"{_cluster_full_name(worst)} 当前使用率 {_percent_label(ratio)}，90 天预测 {_bytes_label((worst.get('forecast') or {}).get('forecast_90d'))}。",
        })
        exhaustion_candidates = [
            (_exhaustion_days(cluster), cluster)
            for cluster in clusters
            if _exhaustion_days(cluster) < float("inf")
        ]
        soonest = min(exhaustion_candidates, key=lambda item: item[0]) if exhaustion_candidates else None
        if soonest:
            days, cluster = soonest
            level = "高" if days <= 90 else "中" if days <= 180 else "低"
            rows.append({
                "item": "预计存储耗尽",
                "status": _days_label(days),
                "level": level,
                "description": f"{_cluster_full_name(cluster)} 按当前趋势预计 {_days_label(days)} 后触达容量上限。",
            })
    if top_vms:
        largest = _largest_vm(top_vms)
        if largest:
            rows.append({
                "item": "重点 VM 容量",
                "status": "需关注" if _vm_current_bytes(largest) >= VM_ALERT_BYTES else "正常",
                "level": "中" if _vm_current_bytes(largest) >= VM_ALERT_BYTES else "低",
                "description": f"{_vm_full_name(largest)} 当前容量 {_bytes_label(_vm_current_bytes(largest))}，建议确认是否存在可清理数据。",
            })
        fastest = top_vms[0]
        rows.append({
            "item": "VM 增长来源",
            "status": "需关注" if float(fastest.get("growth_amount") or 0) > 0 else "健康",
            "level": "中" if float(fastest.get("growth_amount") or 0) > 0 else "低",
            "description": f"{_vm_full_name(fastest)} 增长 {_bytes_label(fastest.get('growth_amount'))}，建议确认业务数据增长合理性。",
        })
    if not rows:
        rows.append({"item": "容量数据", "status": "数据不足", "level": "中", "description": "当前缺少可分析的容量或 VM 增长数据。"})
    return rows


def _customer_operation_advice(clusters: list[dict[str, Any]], top_vms: list[dict[str, Any]], profile: ReportPeriodProfile | None = None) -> list[tuple[str, list[str]]]:
    top_vm = top_vms[0] if top_vms else None
    largest = _largest_vm(top_vms)
    high_clusters = [cluster for cluster in clusters if _cluster_used_ratio(cluster) >= 0.8]
    short_items = [
        (
            f"建议确认 {_vm_scope_name(top_vm)} 下的 {_vm_display_name(top_vm)} "
            f"在统计窗口内增长 {_bytes_label(top_vm.get('growth_amount'))} 的业务来源，判断是否为预期写入、日志膨胀或临时数据堆积。"
        )
        if top_vm else "确认当前采集状态和 Prometheus 历史指标完整性，避免因数据缺口影响容量判断。",
        (
            f"对 {_vm_scope_name(largest)} 下的 {_vm_display_name(largest)}（{_bytes_label(_vm_current_bytes(largest))}）"
            "进行存储空间审计，识别可清理的历史数据、快照或日志。"
        )
        if largest else "检查大容量 VM、快照和备份策略，识别可清理空间。",
    ]
    middle_items = [
        "建立月度容量复盘机制，持续跟踪 Top 10 增长最快的虚拟机。",
        "定期复核大容量 VM 和持续增长 VM 的业务归属、数据保留策略和清理窗口，避免单业务长期占用过多集群空间。",
    ]
    long_items = [
        "当容量使用率达到 60% 时启动扩容评估预案，达到 75% 后进入扩容计划跟踪，达到 80% 后优先执行扩容或清理。",
        "定期审查 VM 快照、备份和日志保留策略，避免冗余数据占用有效存储空间。",
    ]
    if high_clusters:
        short_items.insert(0, f"集群 {_cluster_full_name(high_clusters[0])} 已超过 80% 高风险阈值，建议立即确认扩容周期和可清理空间。")
    short_title = (profile or REPORT_PERIOD_PROFILES[30]).short_advice_title
    return [(short_title, short_items), ("中期（1-3 个月）", middle_items), ("三个月以上", long_items)]


def _risk_text_color(level: str) -> str:
    if level == "高":
        return RATIO_RED
    if level == "中":
        return RATIO_ORANGE
    return "27AE60"


def _risk_word_color(label: str) -> str:
    if "高" in label or "风险" in label and "正常" not in label:
        return RATIO_RED
    if "关注" in label or "中" in label:
        return RATIO_ORANGE
    if "正常" in label:
        return "27AE60"
    return TEXT_DARK


def _set_cell_text_color(cell: Any, color: str) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor.from_string(color)
            run.bold = True


def _vm_display_name(vm: dict[str, Any] | None) -> str:
    if not vm:
        return "未知 VM"
    labels = vm.get("labels") or {}
    return str(labels.get("vm") or labels.get("vm_name") or labels.get("vm_id") or "未知 VM")


def _vm_full_name(vm: dict[str, Any] | None) -> str:
    if not vm:
        return "未知 VM"
    labels = vm.get("labels") or {}
    tower = labels.get("tower") or labels.get("tower_id")
    cluster = labels.get("cluster") or labels.get("cluster_id")
    vm_name = _vm_display_name(vm)
    prefix = " / ".join(str(value) for value in [tower, cluster] if value)
    return f"{prefix} / {vm_name}" if prefix else vm_name


def _vm_scope_name(vm: dict[str, Any] | None) -> str:
    if not vm:
        return "未知范围"
    labels = vm.get("labels") or {}
    tower = labels.get("tower") or labels.get("tower_id")
    cluster = labels.get("cluster") or labels.get("cluster_id")
    scope = " / ".join(str(value) for value in [tower, cluster] if value)
    return scope or "未知范围"


def _vm_current_bytes(vm: dict[str, Any] | None) -> float:
    if not vm:
        return 0.0
    return float((vm.get("forecast") or {}).get("current") or 0.0)


def _largest_vm(vms: list[dict[str, Any]]) -> dict[str, Any] | None:
    if not vms:
        return None
    return max(vms, key=_vm_current_bytes)


def _exhaustion_days(cluster: dict[str, Any]) -> float:
    value = (cluster.get("forecast") or {}).get("exhaustion_days")
    try:
        return float(value)
    except (TypeError, ValueError):
        return float("inf")


def _cluster_period_growth(cluster: dict[str, Any], days: int) -> float:
    points = _cluster_points(cluster)
    forecast = cluster.get("forecast") or {}
    current = _float_or_none(forecast.get("current"))
    if current is None and points:
        current = points[-1][1]
    current = current or 0.0
    if len(points) >= 2:
        target = points[-1][0] - days * 86_400
        candidates = [point for point in points if point[0] <= target]
        baseline = candidates[-1][1] if candidates else points[0][1]
        return max(0.0, current - baseline)
    return max(0.0, float(forecast.get("slope_per_day") or 0) * days)


def _cluster_points(cluster: dict[str, Any]) -> list[tuple[int, float]]:
    points = []
    for point in cluster.get("points") or []:
        try:
            points.append((int(float(point[0])), float(point[1])))
        except (TypeError, ValueError, IndexError):
            continue
    return sorted(points)


def _merged_cluster_points(clusters: list[dict[str, Any]]) -> list[tuple[int, float]]:
    by_ts: dict[int, float] = defaultdict(float)
    for cluster in clusters:
        for ts, value in _cluster_points(cluster):
            by_ts[ts] += value
    return sorted(by_ts.items())


def _vms_by_cluster(vms: list[dict[str, Any]]) -> dict[tuple[str, str], list[dict[str, Any]]]:
    grouped: dict[tuple[str, str], list[dict[str, Any]]] = defaultdict(list)
    for vm in vms:
        grouped[_cluster_key(vm.get("labels", {}))].append(vm)
    return grouped


def _cluster_key(labels: dict[str, Any]) -> tuple[str, str]:
    return (str(labels.get("tower_id") or ""), str(labels.get("cluster_id") or ""))


def _top_vms(vms: list[dict[str, Any]], mode: str) -> list[dict[str, Any]]:
    key = (lambda vm: float(vm.get("growth_ratio") or 0)) if mode == "ratio" else (lambda vm: float(vm.get("growth_amount") or 0))
    return sorted(vms, key=key, reverse=True)[:100]


def _merge_growth_candidates(primary: list[dict[str, Any]], secondary: list[dict[str, Any]]) -> list[dict[str, Any]]:
    merged: list[dict[str, Any]] = []
    seen: set[tuple[str, str, str]] = set()
    for vm in [*primary, *secondary]:
        labels = vm.get("labels") or {}
        key = (
            str(labels.get("tower_id") or ""),
            str(labels.get("cluster_id") or ""),
            str(labels.get("vm_id") or labels.get("vm") or ""),
        )
        if key in seen:
            continue
        seen.add(key)
        merged.append(vm)
    return merged


def _float_or_none(value: Any) -> float | None:
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _first_cluster_name(clusters: list[dict[str, Any]]) -> str | None:
    if not clusters:
        return None
    labels = clusters[0].get("labels", {})
    return labels.get("cluster") or labels.get("cluster_id")


def _is_alert_vm(vm: dict[str, Any]) -> bool:
    return float(vm.get("growth_ratio") or 0) > VM_ALERT_RATIO and float(vm.get("growth_amount") or 0) > VM_ALERT_BYTES


def _shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _docx_bytes(document: Document) -> bytes:
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _load_customer_xlsx_template() -> Workbook:
    if XLSX_TEMPLATE_PATH.exists():
        return load_workbook(XLSX_TEMPLATE_PATH)
    workbook = Workbook()
    workbook.active.title = "封面"
    for name in ["执行摘要", "容量趋势", "VM增长TOP20", "日增长详情"]:
        workbook.create_sheet(name)
    return workbook


def _replace_sheet(workbook: Workbook, name: str):
    if name in workbook.sheetnames:
        index = workbook.sheetnames.index(name)
        workbook.remove(workbook[name])
        return workbook.create_sheet(name, index)
    return workbook.create_sheet(name)


def _remove_sheets(workbook: Workbook, names: list[str]) -> None:
    for name in names:
        if name in workbook.sheetnames and len(workbook.sheetnames) > 1:
            workbook.remove(workbook[name])


def _clear_xlsx_sheet(sheet) -> None:
    for row in sheet.iter_rows():
        for cell in row:
            if not isinstance(cell, MergedCell):
                cell.value = None


def _reset_xlsx_sheet_rows(sheet) -> None:
    for merged_range in list(sheet.merged_cells.ranges):
        sheet.unmerge_cells(str(merged_range))
    if sheet.max_row:
        sheet.delete_rows(1, sheet.max_row)


def _set_xlsx_cell(sheet, coordinate: str, value: Any, *, size: float | None = None, bold: bool | None = None, color: str | None = None, align: str | None = None) -> None:
    cell = sheet[coordinate]
    cell.value = value
    cell.font = Font(
        name="Noto Sans CJK SC",
        size=size if size is not None else cell.font.sz,
        bold=bold if bold is not None else cell.font.bold,
        color=color if color is not None else cell.font.color,
    )
    if align is not None:
        cell.alignment = Alignment(horizontal=align, vertical="center", wrap_text=True)


def _write_xlsx_template_cover(sheet, report: dict[str, Any], context: dict[str, Any], settings: V2Settings) -> None:
    clusters = report.get("clusters") or []
    _clear_xlsx_sheet(sheet)
    _set_xlsx_cell(sheet, "B7", REPORT_PRODUCT_NAME, size=13, color=GROWTH_BLUE, align="center")
    _set_xlsx_cell(sheet, "B8", REPORT_COVER_TITLE, size=26, bold=True, color=ACCENT_DARK, align="center")
    _set_xlsx_cell(sheet, "B9", REPORT_COVER_SUBTITLE, size=12, color=TEXT_MUTED, align="center")
    _set_xlsx_cell(sheet, "B11", "───────────────────────────────────────────────────────", size=7, color="CCCCCC", align="center")
    _set_xlsx_cell(sheet, "B13", "客户名称：", size=11, color="555555", align="center")
    _set_xlsx_cell(sheet, "B14", f"Tower范围：{_tower_scope_label(clusters, context['scope_label'])}", size=11, color="555555", align="center")
    _set_xlsx_cell(sheet, "B15", f"集群范围：{_cluster_scope_label(clusters, context['scope_label'])}", size=11, color="555555", align="center")
    _set_xlsx_cell(sheet, "B16", f"统计窗口：{_period_window_label(report)}", size=11, color="555555", align="center")
    _set_xlsx_cell(sheet, "B17", f"预测窗口：{report.get('forecast_days', 90)} 天", size=11, color="555555", align="center")
    _set_xlsx_cell(sheet, "B20", f"本报告由 {REPORT_PRODUCT_NAME} {settings.app_version} 生成", size=9, color=TEXT_MUTED, align="center")


def _write_xlsx_template_summary(
    sheet,
    report: dict[str, Any],
    context: dict[str, Any],
    clusters: list[dict[str, Any]],
    settings: V2Settings,
    profile: ReportPeriodProfile,
) -> None:
    _clear_xlsx_sheet(sheet)
    current = sum(float((cluster.get("forecast") or {}).get("current") or 0) for cluster in clusters)
    forecast_90d = sum(float((cluster.get("forecast") or {}).get("forecast_90d") or 0) for cluster in clusters)
    total = sum(float(cluster.get("total") or 0) for cluster in clusters)
    growth = sum(_cluster_period_growth(cluster, profile.days) for cluster in clusters)
    risk_status, _, risk_note = _overall_risk_status(clusters)
    growth_ratio = growth / max(current - growth, 1) if growth > 0 else 0.0
    usage_ratio = current / total if total else 0.0
    forecast_ratio = forecast_90d / total if total else 0.0

    _set_xlsx_cell(sheet, "A1", "一、执行摘要", size=16, bold=True, color=ACCENT_DARK)
    _set_xlsx_cell(
        sheet,
        "A2",
        f"报告周期：{_period_window_label(report)}  |  Tower：{_tower_scope_label(clusters, context['scope_label'])}  |  集群：{_cluster_scope_label(clusters, context['scope_label'])}  |  统计口径：{profile.window_growth_label}",
        size=9,
        color=TEXT_MUTED,
    )
    headers = ["当前已用容量", profile.window_growth_label, "90 天预测容量", "风险状态"]
    values = [_xlsx_bytes_label(current), _xlsx_signed_bytes_label(growth), _xlsx_bytes_label(forecast_90d), risk_status]
    subtitles = [f"使用率 {_percent_label(usage_ratio)}", f"增长率 {_percent_label(growth_ratio)}", f"预计使用率 {_percent_label(forecast_ratio)}", risk_note]
    for column, (header, value, subtitle) in enumerate(zip(headers, values, subtitles), start=1):
        for row in [4, 5, 6]:
            cell = sheet.cell(row=row, column=column)
            cell.fill = PatternFill(fill_type="solid", fgColor=ACCENT_LIGHT)
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        sheet.cell(row=4, column=column).value = header
        sheet.cell(row=4, column=column).font = Font(name="Noto Sans CJK SC", size=9, color=TEXT_MUTED)
        sheet.cell(row=5, column=column).value = value
        sheet.cell(row=5, column=column).font = Font(name="Noto Sans CJK SC", size=18, bold=True, color=RATIO_RED if risk_status == "高风险" and column == 4 else ACCENT_DARK)
        sheet.cell(row=6, column=column).value = subtitle
        sheet.cell(row=6, column=column).font = Font(name="Noto Sans CJK SC", size=8, color=TEXT_MUTED)

    _set_xlsx_cell(sheet, "A8", "关键发现", size=13, bold=True, color=ACCENT_DARK)
    findings = _customer_key_findings(clusters, _customer_growth_vms(report), risk_note, context)
    for offset, finding in enumerate(findings[:4], start=9):
        _set_xlsx_cell(sheet, f"A{offset}", finding, size=10, color=TEXT_DARK)
    _set_xlsx_cell(sheet, "A13", "容量风险摘要", size=13, bold=True, color=ACCENT_DARK)
    _set_xlsx_cell(sheet, "A14", _capacity_risk_summary(report), size=10, color=TEXT_DARK)
    _set_xlsx_cell(sheet, "A15", f"当前软件版本：{settings.app_version}", size=9, color=TEXT_MUTED)
    _set_xlsx_cell(sheet, "A16", _profile_sample_notice(report, profile), size=9, color=TEXT_MUTED)
    for row in [1, 2, 8, 9, 10, 11, 12, 13, 14, 15, 16]:
        _merge_title_row(sheet, row, 1, 6)
    for row in [2, 9, 10, 11, 12, 16]:
        sheet.row_dimensions[row].height = 38
    _autosize(sheet)
    sheet.column_dimensions["A"].width = max(sheet.column_dimensions["A"].width or 0, 42)
    for column in ["B", "C", "D", "E", "F"]:
        sheet.column_dimensions[column].width = max(sheet.column_dimensions[column].width or 0, 16)


def _write_xlsx_template_capacity_trend(sheet, report: dict[str, Any], context: dict[str, Any], clusters: list[dict[str, Any]], profile: ReportPeriodProfile) -> None:
    _reset_xlsx_sheet_rows(sheet)
    sheet.append(["二、集群容量趋势"])
    sheet.append([f"统计窗口：{_period_window_label(report)}；{_profile_sample_notice(report, profile)}"])
    headers = ["Tower", "集群", "当前已用容量", profile.window_growth_label, "90 天预测容量", "总容量", "使用率", "预计耗尽天数", "风险状态"]
    sheet.append(headers)
    if not clusters:
        sheet.append(["当前范围暂无集群容量数据"] + [""] * (len(headers) - 1))
    for cluster in clusters:
        labels = cluster.get("labels") or {}
        forecast = cluster.get("forecast") or {}
        risk, _ = _risk_level(cluster)
        sheet.append(
            [
                labels.get("tower") or labels.get("tower_id") or "",
                labels.get("cluster") or labels.get("cluster_id") or "",
                _xlsx_bytes_label(forecast.get("current")),
                _xlsx_signed_bytes_label(_cluster_period_growth(cluster, profile.days)),
                _xlsx_bytes_label(forecast.get("forecast_90d")),
                _xlsx_bytes_label(cluster.get("total")),
                _percent_label(_cluster_used_ratio(cluster)),
                _days_label(forecast.get("exhaustion_days")),
                risk,
            ]
        )
    _style_customer_xlsx_table(sheet, title_rows={1, 2}, header_rows={3})
    sheet.freeze_panes = "A4"
    sheet.auto_filter.ref = f"A3:{get_column_letter(len(headers))}{sheet.max_row}"


def _write_xlsx_template_vm_top100(sheet, vms: list[dict[str, Any]], report: dict[str, Any], profile: ReportPeriodProfile) -> None:
    _reset_xlsx_sheet_rows(sheet)
    window_label = _vm_sample_window_label(vms, report)
    sheet.append([f"三、{profile.vm_growth_title} TOP100（{window_label}）", "", "", "", "", "", "", "", "四、虚拟机增长率 TOP100"])
    sheet.append([])
    sheet.append(["按增长量降序", "", "", "", "", "", "", "", "按增长率降序"])
    amount_headers = ["排名", "虚拟机名称", "当前容量", "期初容量", "增长量", "增长率", "风险"]
    ratio_headers = ["排名", "虚拟机名称", "当前容量", "期初容量", "增长量", "增长率"]
    for index, value in enumerate(amount_headers, start=1):
        sheet.cell(row=4, column=index).value = value
    for index, value in enumerate(ratio_headers, start=9):
        sheet.cell(row=4, column=index).value = value
    amount_vms = _top_vms(vms, "amount")[:100]
    ratio_vms = _top_vms(vms, "ratio")[:100]
    if not amount_vms and not ratio_vms:
        sheet.cell(row=5, column=1).value = profile.vm_empty_text
    for row_index, vm in enumerate(amount_vms, start=5):
        for column, value in enumerate(_xlsx_vm_display_row(vm, row_index - 4, include_risk=True), start=1):
            sheet.cell(row=row_index, column=column).value = value
    for row_index, vm in enumerate(ratio_vms, start=5):
        for column, value in enumerate(_xlsx_vm_display_row(vm, row_index - 4, include_risk=False), start=9):
            sheet.cell(row=row_index, column=column).value = value
    _style_customer_xlsx_table(sheet, title_rows={1, 3}, header_rows={4})
    _apply_vm_top100_layout(sheet, left_header_row=4, right_header_row=4)
    sheet.freeze_panes = "A5"


def _write_xlsx_template_day_growth(sheet, vms: list[dict[str, Any]], report: dict[str, Any]) -> None:
    _reset_xlsx_sheet_rows(sheet)
    sheet.append([f"日增长最快虚拟机（{_period_window_label(report)}）"])
    sheet.append([])
    headers = ["排名", "虚拟机名称", "Tower", "集群", "当前容量", "期初容量", "日增长量", "增长率", "风险"]
    sheet.append(headers)
    if not vms:
        sheet.append(["暂无日增长 VM 数据"])
    for index, vm in enumerate(vms[:100], start=1):
        row = _xlsx_vm_display_row(vm, index, include_risk=True)
        labels = vm.get("labels") or {}
        sheet.append([row[0], row[1], labels.get("tower") or labels.get("tower_id") or "", labels.get("cluster") or labels.get("cluster_id") or "", *row[2:]])
    _style_customer_xlsx_table(sheet, title_rows={1}, header_rows={3})
    sheet.freeze_panes = "A4"
    sheet.auto_filter.ref = f"A3:{get_column_letter(len(headers))}{sheet.max_row}"


def _write_scope_detail_sheet(sheet, clusters: list[dict[str, Any]]) -> None:
    headers = ["序号", "Tower", "集群", "Tower ID", "集群 ID"]
    sheet.append(headers)
    if not clusters:
        sheet.append(["-", "当前范围暂无集群容量数据", "", "", ""])
    for index, cluster in enumerate(clusters, start=1):
        labels = cluster.get("labels") or {}
        sheet.append([
            index,
            labels.get("tower") or labels.get("tower_id") or "",
            labels.get("cluster") or labels.get("cluster_id") or "",
            labels.get("tower_id") or "",
            labels.get("cluster_id") or "",
        ])
    _style_customer_xlsx_table(sheet, header_rows={1})
    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = f"A1:{get_column_letter(len(headers))}{sheet.max_row}"


def _write_cluster_summary_sheet(sheet, clusters: list[dict[str, Any]], report: dict[str, Any], profile: ReportPeriodProfile) -> None:
    sheet.append(["集群汇总"])
    sheet.append([f"统计窗口：{_period_window_label(report)}；{_profile_sample_notice(report, profile)}"])
    headers = ["Tower", "集群", "当前容量", profile.window_growth_label, "90 天预测容量", "总容量", "使用率", "预计耗尽天数", "风险"]
    sheet.append(headers)
    if not clusters:
        sheet.append(["当前范围暂无集群容量数据"] + [""] * (len(headers) - 1))
    for cluster in clusters:
        labels = cluster.get("labels") or {}
        forecast = cluster.get("forecast") or {}
        risk, _ = _risk_level(cluster)
        sheet.append([
            labels.get("tower") or labels.get("tower_id") or "",
            labels.get("cluster") or labels.get("cluster_id") or "",
            _xlsx_bytes_label(forecast.get("current")),
            _xlsx_signed_bytes_label(_cluster_period_growth(cluster, profile.days)),
            _xlsx_bytes_label(forecast.get("forecast_90d")),
            _xlsx_bytes_label(cluster.get("total")),
            _percent_label(_cluster_used_ratio(cluster)),
            _days_label(forecast.get("exhaustion_days")),
            risk,
        ])
    _style_customer_xlsx_table(sheet, title_rows={1, 2}, header_rows={3})
    sheet.freeze_panes = "A4"
    sheet.auto_filter.ref = f"A3:{get_column_letter(len(headers))}{sheet.max_row}"


def _style_customer_xlsx_table(sheet, *, title_rows: set[int] | None = None, header_rows: set[int] | None = None) -> None:
    title_rows = title_rows or set()
    header_rows = header_rows or {1}
    for row in sheet.iter_rows():
        for cell in row:
            if isinstance(cell, MergedCell):
                continue
            cell.alignment = Alignment(vertical="center", wrap_text=True)
            cell.font = Font(name="Noto Sans CJK SC", size=10, color=TEXT_DARK)
            if cell.row in title_rows:
                cell.font = Font(name="Noto Sans CJK SC", size=16 if cell.row == 1 else 11, bold=True, color=ACCENT_DARK)
            if cell.row in header_rows:
                cell.font = Font(name="Noto Sans CJK SC", size=10, bold=True, color="FFFFFF")
                cell.fill = PatternFill(fill_type="solid", fgColor=ACCENT_DARK)
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            elif cell.row not in title_rows and cell.row % 2 == 0:
                cell.fill = PatternFill(fill_type="solid", fgColor=ACCENT_SOFT)
    _autosize(sheet)


def _merge_title_row(sheet, row: int, start_col: int, end_col: int) -> None:
    if end_col <= start_col:
        return
    range_ref = f"{get_column_letter(start_col)}{row}:{get_column_letter(end_col)}{row}"
    if range_ref not in {str(merged) for merged in sheet.merged_cells.ranges}:
        sheet.merge_cells(range_ref)
    cell = sheet.cell(row=row, column=start_col)
    cell.alignment = Alignment(vertical="center", wrap_text=True)


def _apply_vm_top100_layout(sheet, *, left_header_row: int, right_header_row: int | None = None) -> None:
    widths = {
        "A": 10,
        "B": 36,
        "C": 16,
        "D": 16,
        "E": 16,
        "F": 12,
        "G": 12,
        "I": 10,
        "J": 36,
        "K": 16,
        "L": 16,
        "M": 16,
        "N": 12,
    }
    for column, width in widths.items():
        sheet.column_dimensions[column].width = max(sheet.column_dimensions[column].width or 0, width)
    for row_index in range(1, sheet.max_row + 1):
        if row_index in {1, 3} or str(sheet.cell(row=row_index, column=1).value or "").endswith("TOP100"):
            sheet.row_dimensions[row_index].height = 34
        elif row_index in {left_header_row, right_header_row}:
            sheet.row_dimensions[row_index].height = 24
    _merge_title_row(sheet, 1, 1, 7)
    if sheet.cell(row=1, column=9).value:
        _merge_title_row(sheet, 1, 9, 14)
    if sheet.cell(row=3, column=1).value:
        _merge_title_row(sheet, 3, 1, 7)
    if sheet.cell(row=3, column=9).value:
        _merge_title_row(sheet, 3, 9, 14)


def _xlsx_vm_display_row(vm: dict[str, Any], rank: int, *, include_risk: bool) -> list[Any]:
    labels = vm.get("labels") or {}
    forecast = vm.get("forecast") or {}
    values: list[Any] = [
        rank,
        labels.get("vm") or labels.get("vm_name") or labels.get("vm_id") or "",
        _xlsx_bytes_label(forecast.get("current")),
        _xlsx_bytes_label(vm.get("previous_value")),
        _xlsx_signed_bytes_label(vm.get("growth_amount")),
        _percent_label(vm.get("growth_ratio")),
    ]
    if include_risk:
        values.append(_vm_risk_label(vm))
    return values


def _vm_risk_label(vm: dict[str, Any]) -> str:
    if _is_alert_vm(vm):
        return "高"
    ratio = _float_or_none(vm.get("growth_ratio")) or 0.0
    amount = _float_or_none(vm.get("growth_amount")) or 0.0
    if ratio >= 0.1 or amount >= 50 * 1024**3:
        return "中"
    return "低"


def _xlsx_bytes_label(value: Any) -> str:
    try:
        numeric = float(value)
    except (TypeError, ValueError):
        return "-"
    units = ["B", "KiB", "MiB", "GiB", "TiB", "PiB"]
    index = 0
    while abs(numeric) >= 1024 and index < len(units) - 1:
        numeric /= 1024
        index += 1
    return f"{numeric:.2f} {units[index]}"


def _xlsx_signed_bytes_label(value: Any) -> str:
    try:
        numeric = float(value)
    except (TypeError, ValueError):
        return "-"
    label = _xlsx_bytes_label(abs(numeric))
    if numeric > 0:
        return f"+{label}"
    if numeric < 0:
        return f"-{label}"
    return label


def _percent_label(value: Any) -> str:
    try:
        return f"{float(value) * 100:.1f}%"
    except (TypeError, ValueError):
        return "-"


def _style_sheet(sheet, header_rows: set[int] | None = None, bytes_cols: set[int] | None = None, percent_cols: set[int] | None = None) -> None:
    header_rows = header_rows or {1}
    bytes_cols = bytes_cols or set()
    percent_cols = percent_cols or set()
    title_fill = PatternFill(fill_type="solid", fgColor=ACCENT)
    header_fill = PatternFill(fill_type="solid", fgColor=ACCENT_LIGHT)
    for row in sheet.iter_rows():
        for cell in row:
            cell.alignment = Alignment(vertical="center", wrap_text=True)
            if cell.row == 1:
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = title_fill
            elif cell.row in header_rows or cell.value in {"集群", "VM", "当前容量"}:
                cell.font = Font(bold=True, color=ACCENT_DARK)
                cell.fill = header_fill
            if cell.column in bytes_cols and isinstance(cell.value, (int, float)):
                cell.number_format = '#,##0'
            if cell.column in percent_cols and isinstance(cell.value, (int, float)):
                cell.number_format = "0.00%"
    _autosize(sheet)
    sheet.freeze_panes = "A2"


def _autosize(sheet) -> None:
    for column_cells in sheet.columns:
        width = max(len(str(cell.value or "")) for cell in column_cells) + 2
        sheet.column_dimensions[get_column_letter(column_cells[0].column)].width = min(max(width, 10), 32)


def _write_directory_sheet(sheet, clusters: list[dict[str, Any]]) -> None:
    sheet.append(["集群目录"])
    sheet.append(["序号", "集群", "Sheet"])
    if not clusters:
        sheet.append(["-", "当前范围暂无集群容量数据", ""])
    for index, cluster in enumerate(clusters, start=1):
        name = _safe_sheet_name(_cluster_name(cluster))
        sheet.append([index, _cluster_name(cluster), name])
        sheet.cell(row=sheet.max_row, column=3).hyperlink = f"#'{name}'!A1"
        sheet.cell(row=sheet.max_row, column=3).style = "Hyperlink"
    _style_sheet(sheet, header_rows={1, 2})


def _write_vm_top_sheet(sheet, vms: list[dict[str, Any]], report: dict[str, Any], profile: ReportPeriodProfile) -> None:
    headers = ["Tower", "集群", "VM", "当前容量", "期初容量", "增长量", "增长率"]
    window_label = _vm_sample_window_label(vms, report)
    sheet.append([f"{profile.vm_growth_title} 增长量 TOP100（按增长量降序，统计窗口：{window_label}）"])
    sheet.append([profile.window_growth_label])
    sheet.append([_profile_sample_notice(report, profile)])
    sheet.append(headers)
    amount_start = sheet.max_row + 1
    if not vms:
        sheet.append([profile.vm_empty_text])
    else:
        for vm in _top_vms(vms, "amount"):
            sheet.append(_vm_xlsx_row(vm, include_cluster=True))
    amount_end = sheet.max_row
    sheet.append([])
    sheet.append([f"{profile.vm_growth_title} 增长率 TOP100（按增长率降序，统计窗口：{window_label}）"])
    sheet.append(headers)
    ratio_start = sheet.max_row + 1
    if not vms:
        sheet.append([profile.vm_empty_text])
    else:
        for vm in _top_vms(vms, "ratio"):
            sheet.append(_vm_xlsx_row(vm, include_cluster=True))
    ratio_end = sheet.max_row
    _style_sheet(sheet, header_rows={1, 2, 3, 4, ratio_start - 2, ratio_start - 1}, bytes_cols={4, 5, 6}, percent_cols={7})
    _style_vm_rows(sheet, amount_start, amount_end, 7)
    _style_vm_rows(sheet, ratio_start, ratio_end, 7)
    _add_excel_table(sheet, "VmAmountSummary", 4, amount_end, len(headers))
    _add_excel_table(sheet, "VmRatioSummary", ratio_start - 1, ratio_end, len(headers))


def _write_simple_vm_sheet(sheet, vms: list[dict[str, Any]], empty_text: str, *, include_growth: bool) -> None:
    headers = ["Tower", "集群", "VM", "当前容量"]
    if include_growth:
        headers.extend(["期初容量", "增长量", "增长率"])
    else:
        headers.append("首次出现时间")
    sheet.append([sheet.title])
    sheet.append(headers)
    if not vms:
        sheet.append([empty_text])
    for vm in vms[:100]:
        labels = vm.get("labels", {})
        row = [
            labels.get("tower") or labels.get("tower_id") or "",
            labels.get("cluster") or labels.get("cluster_id") or "",
            labels.get("vm") or labels.get("vm_id") or "",
            _xlsx_bytes_label((vm.get("forecast") or {}).get("current")),
        ]
        if include_growth:
            row.extend([
                _xlsx_bytes_label(vm.get("previous_value")),
                _xlsx_signed_bytes_label(vm.get("growth_amount")),
                _percent_label(vm.get("growth_ratio")),
            ])
        else:
            row.append(vm.get("first_seen_at") or "")
        sheet.append(row)
    _style_customer_xlsx_table(sheet, title_rows={1}, header_rows={2})
    sheet.freeze_panes = "A3"
    sheet.auto_filter.ref = f"A2:{get_column_letter(len(headers))}{sheet.max_row}"
    for column, width in {"A": 16, "B": 18, "C": 36, "D": 16, "E": 16, "F": 16, "G": 12}.items():
        sheet.column_dimensions[column].width = max(sheet.column_dimensions[column].width or 0, width)


def _write_cluster_vm_sheet(sheet, cluster: dict[str, Any], vms: list[dict[str, Any]], report: dict[str, Any], profile: ReportPeriodProfile) -> None:
    _reset_xlsx_sheet_rows(sheet)
    forecast = cluster.get("forecast") or {}
    risk, _ = _risk_level(cluster)
    labels = cluster.get("labels") or {}
    sheet.append([_cluster_full_name(cluster)])
    sheet.append(
        [
            "Tower",
            labels.get("tower") or labels.get("tower_id") or "",
            "集群",
            labels.get("cluster") or labels.get("cluster_id") or "",
            "当前容量",
            _xlsx_bytes_label(forecast.get("current")),
            profile.window_growth_label,
            _xlsx_signed_bytes_label(_cluster_period_growth(cluster, profile.days)),
            "风险",
            risk,
        ]
    )
    sheet.append(
        [
            "90 天预测容量",
            _xlsx_bytes_label(forecast.get("forecast_90d")),
            "总容量",
            _xlsx_bytes_label(cluster.get("total")),
            "使用率",
            _percent_label(_cluster_used_ratio(cluster)),
            "预计耗尽天数",
            _days_label(forecast.get("exhaustion_days")),
        ]
    )
    sheet.append([_profile_sample_notice(report, profile)])
    sheet.append([])
    headers = ["VM", "当前容量", "期初容量", "增长量", "增长率"]
    window_label = _vm_sample_window_label(vms, report)
    sheet.append([f"{profile.vm_growth_title} 增长量 TOP100（按增长量降序，统计窗口：{window_label}）"])
    sheet.append(headers)
    amount_start = sheet.max_row + 1
    if not vms:
        sheet.append([profile.vm_empty_text])
    else:
        for vm in _top_vms(vms, "amount"):
            sheet.append(_vm_xlsx_row(vm, include_cluster=False))
    amount_end = sheet.max_row
    sheet.append([])
    sheet.append([f"{profile.vm_growth_title} 增长率 TOP100（按增长率降序，统计窗口：{window_label}）"])
    sheet.append(headers)
    ratio_start = sheet.max_row + 1
    if not vms:
        sheet.append([profile.vm_empty_text])
    else:
        for vm in _top_vms(vms, "ratio"):
            sheet.append(_vm_xlsx_row(vm, include_cluster=False))
    ratio_end = sheet.max_row
    _style_customer_xlsx_table(sheet, title_rows={1, 4, 6, ratio_start - 1}, header_rows={2, 3, 7, ratio_start})
    _style_vm_rows(sheet, amount_start, amount_end, 5)
    _style_vm_rows(sheet, ratio_start, ratio_end, 5)
    _merge_title_row(sheet, 1, 1, 10)
    _merge_title_row(sheet, 4, 1, 10)
    _merge_title_row(sheet, 6, 1, 5)
    _merge_title_row(sheet, ratio_start - 1, 1, 5)
    _apply_cluster_sheet_layout(sheet, amount_header_row=7, ratio_header_row=ratio_start)
    sheet.freeze_panes = "A7"
    _add_excel_table(sheet, _table_safe_name(sheet.title, "Amount"), 7, amount_end, len(headers))
    _add_excel_table(sheet, _table_safe_name(sheet.title, "Ratio"), ratio_start, ratio_end, len(headers))


def _apply_cluster_sheet_layout(sheet, *, amount_header_row: int, ratio_header_row: int) -> None:
    widths = {
        "A": 36,
        "B": 16,
        "C": 16,
        "D": 16,
        "E": 12,
        "F": 16,
        "G": 16,
        "H": 16,
        "I": 12,
        "J": 12,
    }
    for column, width in widths.items():
        sheet.column_dimensions[column].width = max(sheet.column_dimensions[column].width or 0, width)
    for row_index in [1, 4, 6, ratio_header_row - 1]:
        sheet.row_dimensions[row_index].height = 32
    for row_index in [2, 3, amount_header_row, ratio_header_row]:
        sheet.row_dimensions[row_index].height = 24


def _vm_xlsx_row(vm: dict[str, Any], *, include_cluster: bool) -> list[Any]:
    labels = vm.get("labels", {})
    forecast = vm.get("forecast", {})
    row: list[Any] = []
    if include_cluster:
        row.extend([labels.get("tower") or labels.get("tower_id") or "", labels.get("cluster") or labels.get("cluster_id") or ""])
    row.extend([
        labels.get("vm") or labels.get("vm_name") or labels.get("vm_id") or "",
        _xlsx_bytes_label(forecast.get("current")),
        _xlsx_bytes_label(vm.get("previous_value")),
        _xlsx_signed_bytes_label(vm.get("growth_amount")),
        _percent_label(vm.get("growth_ratio")),
    ])
    return row


def _style_vm_rows(sheet, start_row: int, end_row: int, column_count: int) -> None:
    if end_row < start_row:
        return
    fill = PatternFill(fill_type="solid", fgColor=VM_ALERT_FILL)
    for row_index in range(start_row, end_row + 1):
        values = [sheet.cell(row=row_index, column=column).value for column in range(1, column_count + 1)]
        amount = _parse_xlsx_bytes_label(values[-2]) if len(values) >= 2 else 0
        ratio = _parse_percent_label(values[-1]) if values else 0
        if _is_alert_vm({"growth_amount": amount, "growth_ratio": ratio}):
            for column in range(1, column_count + 1):
                sheet.cell(row=row_index, column=column).fill = fill


def _parse_xlsx_bytes_label(value: Any) -> float:
    if isinstance(value, (int, float)):
        return float(value)
    text = str(value or "").strip().replace(",", "")
    sign = -1.0 if text.startswith("-") else 1.0
    text = text.lstrip("+-").strip()
    parts = text.split()
    if not parts:
        return 0.0
    try:
        number = float(parts[0])
    except (TypeError, ValueError):
        return 0.0
    unit = parts[1] if len(parts) > 1 else "B"
    multipliers = {
        "B": 1,
        "KiB": 1024,
        "MiB": 1024**2,
        "GiB": 1024**3,
        "TiB": 1024**4,
        "PiB": 1024**5,
        "KB": 1024,
        "MB": 1024**2,
        "GB": 1024**3,
        "TB": 1024**4,
        "PB": 1024**5,
    }
    return sign * number * multipliers.get(unit, 1)


def _parse_percent_label(value: Any) -> float:
    if isinstance(value, (int, float)):
        return float(value)
    text = str(value or "").strip()
    if not text.endswith("%"):
        return 0.0
    try:
        return float(text[:-1]) / 100
    except (TypeError, ValueError):
        return 0.0


def _add_excel_table(sheet, name: str, header_row: int, end_row: int, column_count: int) -> None:
    if end_row <= header_row:
        return
    ref = f"A{header_row}:{get_column_letter(column_count)}{end_row}"
    table = Table(displayName=name[:255], ref=ref)
    table.tableStyleInfo = TableStyleInfo(name="TableStyleMedium2", showRowStripes=True, showColumnStripes=False)
    sheet.add_table(table)


def _add_xlsx_growth_chart(sheet, start_row: int, end_row: int) -> None:
    if end_row <= start_row:
        return
    chart = BarChart()
    chart.type = "bar"
    chart.style = 10
    chart.title = "集群容量增长对比"
    chart.y_axis.title = "容量增长"
    chart.x_axis.title = "集群"
    data = Reference(sheet, min_col=2, max_col=4, min_row=start_row, max_row=end_row)
    cats = Reference(sheet, min_col=1, min_row=start_row + 1, max_row=end_row)
    chart.add_data(data, titles_from_data=True)
    chart.set_categories(cats)
    chart.height = 8
    chart.width = 18
    sheet.add_chart(chart, "L2")


def _safe_sheet_name(value: str) -> str:
    cleaned = "".join(char if char not in "[]:*?/\\'" else "_" for char in str(value))
    return (cleaned[:31] or "Sheet").strip()


def _table_safe_name(sheet_name: str, suffix: str) -> str:
    return "".join(char for char in f"{sheet_name}_{suffix}" if char.isalnum())[:200] or f"Table{suffix}"


def _scope_trend_line_chart(clusters: list[dict[str, Any]], title: str) -> BytesIO | None:
    return _line_chart_image(_merged_cluster_points(clusters), title=title, ylabel="容量 (TiB)")


def _cluster_trend_line_chart(cluster: dict[str, Any], title: str) -> BytesIO | None:
    return _line_chart_image(_cluster_points(cluster), title=title, ylabel="容量 (TiB)")


def _cluster_top_growth_bar_chart(clusters: list[dict[str, Any]], title: str) -> BytesIO | None:
    items = [(_cluster_name(cluster), _cluster_period_growth(cluster, 30)) for cluster in clusters if _cluster_period_growth(cluster, 30) > 0]
    return _horizontal_bar_chart_image(items, title=title, unit="TiB", limit=5, scale="tib")


def _vm_top_growth_bar_chart(vms: list[dict[str, Any]], title: str) -> BytesIO | None:
    items = []
    for vm in _top_vms(vms, "amount")[:10]:
        labels = vm.get("labels", {})
        value = float(vm.get("growth_amount") or 0)
        if value:
            items.append((str(labels.get("vm") or labels.get("vm_id") or "VM"), value))
    return _horizontal_bar_chart_image(items, title=title, unit="GB", limit=10, scale="gb")


def _line_chart_image(points: list[tuple[int, float]], title: str, ylabel: str) -> BytesIO | None:
    if len(points) < 2:
        return None
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.dates as mdates
        import matplotlib.pyplot as plt
        _configure_chart_fonts(matplotlib)
        matplotlib.rcParams["font.family"] = [CHART_FONT_FAMILY, "DejaVu Serif"]
        matplotlib.rcParams["axes.unicode_minus"] = False
    except Exception:
        return None
    dates = [datetime.fromtimestamp(ts) for ts, _ in points]
    values = [_bytes_to_tib(value) for _, value in points]
    fig, ax = plt.subplots(figsize=(6.6, 4.0), dpi=180)
    ax.plot(dates, values, color="#003BFF", linewidth=1.4)
    ax.set_title(title, fontsize=12, pad=8)
    ax.set_xlabel("时间", fontsize=10)
    ax.set_ylabel(ylabel, fontsize=10)
    y_min, y_max = _chart_y_limits(values)
    ax.set_ylim(y_min, y_max)
    ax.grid(axis="y", color="#E6EDF5", linewidth=0.6)
    ax.xaxis.set_major_locator(mdates.AutoDateLocator(minticks=4, maxticks=6))
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%Y-%m-%d"))
    ax.tick_params(axis="both", labelsize=9, colors="#333333")
    fig.tight_layout()
    return _figure_bytes(fig, plt)


def _chart_color(value: str) -> str:
    color = value.strip()
    return color if color.startswith("#") else f"#{color}"


def _horizontal_bar_chart_image(items: list[tuple[str, float]], title: str, unit: str, limit: int, scale: str) -> BytesIO | None:
    if not items:
        return None
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        _configure_chart_fonts(matplotlib)
        matplotlib.rcParams["font.family"] = [CHART_FONT_FAMILY, "DejaVu Serif"]
        matplotlib.rcParams["axes.unicode_minus"] = False
    except Exception:
        return None
    ranked = sorted(items, key=lambda item: item[1], reverse=True)[:limit]
    names = [_truncate_label(name) for name, _ in ranked][::-1]
    values = [_chart_bar_value(value, scale) for _, value in ranked][::-1]
    max_value = max(values) if values else 0
    fig, ax = plt.subplots(figsize=(6.8, max(2.2, 0.45 * len(values) + 1.1)), dpi=180)
    bars = ax.barh(names, values, color="#1155CC", height=0.32)
    ax.set_title(title, fontsize=12, pad=8)
    ax.set_xlim(0, max(max_value * 1.18, 1))
    ax.grid(axis="x", color="#E6EDF5", linewidth=0.6)
    for bar, value in zip(bars, values):
        ax.text(bar.get_width(), bar.get_y() + bar.get_height() / 2, f"{value:.2f} {unit}", va="center", ha="left", fontsize=9)
    fig.tight_layout()
    return _figure_bytes(fig, plt)


def _configure_chart_fonts(matplotlib: Any) -> None:
    font_path = Path(CHART_FONT_PATH)
    if not font_path.exists():
        return
    try:
        from matplotlib import font_manager
        font_manager.fontManager.addfont(str(font_path))
    except Exception:
        return


def _figure_bytes(fig: Any, plt: Any) -> BytesIO:
    output = BytesIO()
    fig.savefig(output, format="png", bbox_inches="tight", facecolor="white")
    plt.close(fig)
    output.seek(0)
    return output


def _chart_y_limits(values: list[float]) -> tuple[float, float]:
    if not values:
        return 0.0, 1.0
    y_min = min(values)
    y_max = max(values)
    padding = max((y_max - y_min) * 0.18, y_max * 0.01, 0.1) if y_min != y_max else max(abs(y_max) * 0.05, 0.1)
    lower = max(0.0, y_min - padding)
    upper = y_max + padding
    return lower, upper if upper > lower else lower + 1.0


def _truncate_label(value: str, limit: int = 22) -> str:
    text = str(value)
    return text if len(text) <= limit else f"{text[:8]}...{text[-8:]}"


def _bytes_to_tib(value: Any) -> float:
    try:
        return float(value or 0) / 1024**4
    except (TypeError, ValueError):
        return 0.0


def _bytes_to_gb(value: Any) -> float:
    try:
        return float(value or 0) / 1024**3
    except (TypeError, ValueError):
        return 0.0


def _chart_bar_value(value: float, scale: str) -> float:
    if scale == "gb":
        return _bytes_to_gb(value)
    return _bytes_to_tib(value)


def _local_now(settings: V2Settings) -> datetime:
    try:
        return datetime.now(ZoneInfo(settings.timezone))
    except ZoneInfoNotFoundError:
        return datetime.now()


def _slug(value: str) -> str:
    normalized = "".join(char if char.isalnum() or char in {"-", "_"} else "-" for char in value.strip())
    return normalized.strip("-") or "all"


def _bytes_label(value: Any) -> str:
    try:
        numeric = float(value)
    except (TypeError, ValueError):
        return "-"
    units = ["B", "KB", "MB", "GB", "TB", "PB"]
    index = 0
    while abs(numeric) >= 1024 and index < len(units) - 1:
        numeric /= 1024
        index += 1
    return f"{numeric:.2f} {units[index]}"


def _signed_bytes_label(value: Any) -> str:
    try:
        numeric = float(value)
    except (TypeError, ValueError):
        return "-"
    label = _bytes_label(abs(numeric))
    if numeric > 0:
        return f"+{label}"
    if numeric < 0:
        return f"-{label}"
    return label


def _days_label(value: Any) -> str:
    try:
        return f"{float(value):.0f} 天"
    except (TypeError, ValueError):
        return "未触发"


def _percent_label(value: Any) -> str:
    try:
        return f"{float(value) * 100:.2f}%"
    except (TypeError, ValueError):
        return "-"


# v2 keeps its report data service, task flow and Excel workbook, but the Word
# document intentionally uses a compact customer-facing delivery template.
def build_report_docx(report: dict[str, Any], settings: V2Settings, *, period_days: int) -> tuple[bytes, str, Path, str]:
    context = _export_context(report, settings, period_days, "docx")
    clusters = report.get("clusters") or []
    month_vms = report.get("month_fastest_growing_vms") or []
    window_vms = report.get("window_fastest_growing_vms") or []
    day_vms = report.get("day_fastest_growing_vms") or []
    growth_vms = _merge_growth_candidates(_customer_growth_vms(report), day_vms)
    document = Document()

    _customer_setup_document(document)
    _customer_add_cover(document, context, settings)
    document.add_page_break()
    _customer_add_native_toc(document)
    document.add_page_break()
    _customer_add_executive_summary(document, context, clusters, growth_vms)
    document.add_page_break()
    _customer_add_cluster_overview(document, context, clusters, _report_cluster_vm_counts(report))
    document.add_page_break()
    _customer_add_vm_growth_analysis(document, context, clusters, growth_vms)
    document.add_page_break()
    _customer_add_chart_section(document, context, clusters, growth_vms)
    document.add_page_break()
    _customer_add_risk_and_advice(document, context, clusters, growth_vms)

    content = _docx_bytes(document)
    return _persist_report(content, settings, context["filename"])


def _customer_setup_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    _v1_clear_paragraph(section.header.paragraphs[0])
    _v1_clear_paragraph(section.footer.paragraphs[0])
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = document.styles[style_name]
        _v1_apply_style_font(style)
    normal = document.styles["Normal"]
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT_DARK)
    for name, size in [("Heading 1", 18), ("Heading 2", 13), ("Heading 3", 11)]:
        style = document.styles[name]
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(ACCENT_DARK)
    _customer_enable_update_fields(document)


def _customer_add_cover(document: Document, context: dict[str, Any], settings: V2Settings) -> None:
    for _ in range(6):
        document.add_paragraph()
    _customer_centered_text(document, REPORT_PRODUCT_NAME, 16, GROWTH_BLUE)
    _customer_centered_text(document, REPORT_COVER_TITLE, 28, ACCENT_DARK, bold=True, before=10)
    _customer_centered_text(document, REPORT_COVER_SUBTITLE, 13, TEXT_MUTED, before=8)
    document.add_paragraph()
    _customer_center_rule(document, width=4100)
    for _ in range(2):
        document.add_paragraph()
    meta = _customer_cover_meta(context, settings)
    for label, value in meta:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(10)
        label_run = paragraph.add_run(f"{label}：")
        _v1_apply_run_font(label_run)
        label_run.font.size = Pt(11)
        label_run.font.color.rgb = RGBColor.from_string(TEXT_MUTED)
        value_run = paragraph.add_run(value)
        _v1_apply_run_font(value_run)
        value_run.bold = True
        value_run.font.size = Pt(11)
        value_run.font.color.rgb = RGBColor.from_string(TEXT_DARK)
    for _ in range(2):
        document.add_paragraph()
    _customer_centered_text(document, f"本报告由 {REPORT_PRODUCT_NAME} {settings.app_version} 自动生成", 9, TEXT_MUTED)


def _customer_add_native_toc(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(14)
    run = paragraph.add_run("目录")
    _v1_apply_run_font(run)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string(ACCENT_DARK)
    _customer_left_rule(paragraph)

    toc_paragraph = document.add_paragraph()
    toc_paragraph.paragraph_format.space_after = Pt(8)
    _customer_append_toc_field(toc_paragraph)


def _customer_append_toc_field(paragraph: Any) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = 'TOC \\o "1-3" \\h \\z \\u'

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "请在 Word/WPS 中更新目录"

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    for element in [begin, instruction, separate, placeholder, end]:
        run = OxmlElement("w:r")
        run.append(element)
        paragraph._p.append(run)


def _customer_enable_update_fields(document: Document) -> None:
    settings = document.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def _customer_add_executive_summary(document: Document, context: dict[str, Any], clusters: list[dict[str, Any]], growth_vms: list[dict[str, Any]]) -> None:
    _customer_section_title(document, "一", "摘要")
    summary = _customer_summary_text(context, clusters)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(18)
    _customer_add_emphasis_text(paragraph, summary, base_size=11)

    _customer_kpi_strip(document, context, clusters)
    _customer_table_note(document, context["profile"].focus_note)
    _customer_subtitle(document, "关键发现")
    for finding in _customer_key_findings(clusters, growth_vms, _overall_risk_status(clusters)[2], context):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.24)
        paragraph.paragraph_format.space_after = Pt(7)
        _customer_add_emphasis_text(paragraph, finding)


def _customer_add_cluster_overview(document: Document, context: dict[str, Any], clusters: list[dict[str, Any]], vm_counts: dict[tuple[str, str], int]) -> None:
    _customer_section_title(document, "二", "集群容量概览")
    primary = clusters[0] if clusters else {}
    _customer_subtitle(document, "2.1  范围基本信息")
    _customer_cluster_info_table(document, context, primary, clusters, vm_counts)
    document.add_paragraph()
    _customer_subtitle(document, "2.2  容量增长趋势")
    _customer_cluster_growth_table(document, context, clusters)
    document.add_paragraph()
    _customer_subtitle(document, "2.3  容量使用率可视化")
    _customer_usage_bars(document, clusters)


def _customer_add_vm_growth_analysis(document: Document, context: dict[str, Any], clusters: list[dict[str, Any]], growth_vms: list[dict[str, Any]]) -> None:
    _customer_section_title(document, "三", "集群虚拟机增长分析")
    profile = context["profile"]
    grouped = _vms_by_cluster(growth_vms)
    if not clusters:
        _customer_table_note(document, "当前范围暂无集群，无法按集群生成 VM 增长分析。")
        return
    for index, cluster in enumerate(clusters, start=1):
        if index > 1:
            document.add_page_break()
        labels = cluster.get("labels") or {}
        cluster_vms = grouped.get(_cluster_key(labels), [])
        _customer_subtitle(document, f"三.{index}  {_cluster_full_name(cluster)}")
        _customer_subtitle(document, f"{profile.vm_growth_title}（增长量 Top 20）", level=3)
        top_amount_vms = _top_vms(cluster_vms, "amount")[:20]
        _customer_vm_window_note(document, context, top_amount_vms, "单位：GB")
        _customer_vm_table(document, top_amount_vms, "amount", empty_text=profile.vm_empty_text)
        _customer_table_note(document, f"注：以上为该集群增长量 Top 20，完整清单共 {len(cluster_vms)} 台 VM。")
        _customer_subtitle(document, f"{profile.vm_growth_title}（增长率 Top 20）", level=3)
        top_ratio_vms = _top_vms(cluster_vms, "ratio")[:20]
        _customer_vm_window_note(document, context, top_ratio_vms, "按增长率降序排列")
        _customer_vm_table(document, top_ratio_vms, "ratio", empty_text=profile.vm_empty_text)


def _customer_add_chart_section(document: Document, context: dict[str, Any], clusters: list[dict[str, Any]], growth_vms: list[dict[str, Any]]) -> None:
    _customer_section_title(document, "四", "集群容量趋势图表")
    profile = context["profile"]
    grouped = _vms_by_cluster(growth_vms)
    if not clusters:
        _customer_table_note(document, "当前范围暂无集群，无法生成集群容量趋势图表。")
        return
    for index, cluster in enumerate(clusters, start=1):
        if index > 1:
            document.add_page_break()
        labels = cluster.get("labels") or {}
        cluster_vms = grouped.get(_cluster_key(labels), [])
        _customer_subtitle(document, f"四.{index}  {_cluster_full_name(cluster)}")
        trend_chart = _cluster_trend_line_chart(cluster, f"{_cluster_full_name(cluster)} 容量使用趋势")
        if trend_chart is not None:
            _customer_add_figure(document, trend_chart, f"图：{_cluster_full_name(cluster)} 容量使用趋势", width=5.9)
        else:
            _customer_table_note(document, "暂无足够历史数据生成该集群容量使用趋势图。")
        top_chart = _vm_top_growth_bar_chart(_top_vms(cluster_vms, "amount")[:10], f"{profile.vm_growth_title} Top 10 VM 增长量")
        if top_chart is not None:
            _customer_add_figure(document, top_chart, f"图：{_cluster_full_name(cluster)} {profile.vm_growth_title} Top 10 VM 增长量", width=6.1)
        else:
            _customer_table_note(document, f"暂无足够 VM 增长数据生成该集群 {profile.vm_growth_title} Top 10 VM 增长量图表。")


def _customer_add_risk_and_advice(document: Document, context: dict[str, Any], clusters: list[dict[str, Any]], growth_vms: list[dict[str, Any]]) -> None:
    _customer_section_title(document, "五", "全集群汇总报告")
    _customer_subtitle(document, "5.1  容量风险评估矩阵")
    _customer_risk_matrix_table(document, clusters, growth_vms)
    document.add_paragraph()
    _customer_subtitle(document, "5.2  运维建议")
    for title, items in _customer_operation_advice(clusters, growth_vms, context["profile"]):
        _customer_advice_title(document, title)
        for item in items:
            body = document.add_paragraph()
            body.paragraph_format.left_indent = Inches(0.24)
            body.paragraph_format.space_after = Pt(5)
            _customer_add_emphasis_text(body, item, base_size=10)
    _customer_table_note(document, "声明：本报告基于平台采集容量数据自动生成，预测结果基于历史增长趋势推算，仅供容量规划参考。")


def _customer_advice_title(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.keep_with_next = True
    # A lightweight marker keeps these compact recommendation group titles
    # visually distinct from body text in Word and LibreOffice renderers.
    run = paragraph.add_run(f"· {title}")
    _v1_apply_run_font(run)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(GROWTH_BLUE)


def _customer_centered_text(document: Document, text: str, size: int, color: str, *, bold: bool = False, before: int = 0) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    _v1_apply_run_font(run)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def _customer_center_rule(document: Document, *, width: int) -> None:
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    side = max(1, int((8200 - width) / 2))
    _v1_set_table_width(table, [side, width, side])
    _v1_set_table_borders(table, "FFFFFF")
    for index, cell in enumerate(table.rows[0].cells):
        _shade_cell(cell, GROWTH_BLUE if index == 1 else "FFFFFF")
        _customer_set_cell_height(cell, 28)
        for margin in ["top", "bottom", "start", "end"]:
            _v1_set_cell_margin(cell, margin, 0)


def _customer_set_cell_height(cell: Any, height: int) -> None:
    tr_pr = cell._tc.getparent().get_or_add_trPr()
    tr_height = tr_pr.find(qn("w:trHeight"))
    if tr_height is None:
        tr_height = OxmlElement("w:trHeight")
        tr_pr.append(tr_height)
    tr_height.set(qn("w:val"), str(height))
    tr_height.set(qn("w:hRule"), "exact")


def _customer_section_title(document: Document, order: str, title: str) -> None:
    paragraph = document.add_paragraph(style="Heading 1")
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(14)
    run = paragraph.add_run(f"{order}  {title}")
    _v1_apply_run_font(run)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string(ACCENT_DARK)
    _customer_left_rule(paragraph)


def _customer_left_rule(paragraph: Any) -> None:
    paragraph.paragraph_format.space_after = Pt(12)
    p_pr = paragraph._p.get_or_add_pPr()
    border = p_pr.find(qn("w:pBdr"))
    if border is None:
        border = OxmlElement("w:pBdr")
        p_pr.append(border)
    bottom = border.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        border.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), GROWTH_BLUE)


def _customer_subtitle(document: Document, text: str, *, level: int = 2) -> None:
    style_name = "Heading 3" if level >= 3 else "Heading 2"
    paragraph = document.add_paragraph(style=style_name)
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    _v1_apply_run_font(run)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor.from_string(ACCENT_DARK)


def _customer_cover_meta(context: dict[str, Any], settings: V2Settings) -> list[tuple[str, str]]:
    clusters = context["clusters"]
    generated_date = str(context["generated_at"]).split(" ")[0]
    return [
        ("客户名称", ""),
        ("Tower范围", _tower_scope_label(clusters, context["scope_label"]).upper()),
        ("集群范围", _cluster_scope_label(clusters, context["scope_label"]).upper()),
        ("统计窗口", _period_window_label(context["report"])),
        ("本报表统计窗口", _requested_report_window_label(context["report"])),
        ("预测窗口", f"{context['report'].get('forecast_days') or 90} 天"),
        ("生成日期", generated_date),
    ]


def _customer_summary_text(context: dict[str, Any], clusters: list[dict[str, Any]]) -> str:
    clusters_label = "、".join(_cluster_full_name(cluster) for cluster in clusters[:3]) or context["scope_label"]
    if len(clusters) > 3:
        clusters_label += f" 等 {len(clusters)} 个集群"
    risk = _overall_risk_status(clusters)[0]
    return (
        f"本报告基于 {REPORT_PRODUCT_NAME} 在 {_period_window_label(context['report'])} "
        f"期间采集的容量数据，对集群 {clusters_label} 的存储使用情况进行分析与趋势预测。"
        f"当前容量风险状态为{risk}。"
    )


def _customer_kpi_strip(document: Document, context: dict[str, Any], clusters: list[dict[str, Any]]) -> None:
    profile = context["profile"]
    total_current = sum(float((cluster.get("forecast") or {}).get("current") or 0) for cluster in clusters)
    total_window = sum(_cluster_period_growth(cluster, _customer_window_days(context["report"])) for cluster in clusters)
    total_forecast_90 = sum(float((cluster.get("forecast") or {}).get("forecast_90d") or 0) for cluster in clusters)
    total_capacity = sum(float(cluster.get("total") or 0) for cluster in clusters)
    risk_label = _overall_risk_status(clusters)[0]
    risk_note = "无容量风险集群" if risk_label == "正常" else "存在容量风险集群"
    values = [
        ("当前已用容量", _bytes_label(total_current), f"总容量 {_bytes_label(total_capacity)}\n使用率 {_percent_label(total_current / total_capacity) if total_capacity > 0 else '-'}"),
        (profile.window_growth_label, _signed_bytes_label(total_window), _period_window_label(context["report"])),
        ("90 天预测容量", _bytes_label(total_forecast_90), f"使用率 {_percent_label(total_forecast_90 / total_capacity) if total_capacity > 0 else '-'}"),
        ("风险状态", risk_label, risk_note),
    ]
    table = document.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    _v1_set_table_width(table, [2050, 2050, 2050, 2050])
    _v1_set_table_borders(table, "FFFFFF")
    for cell, (label, value, note) in zip(table.rows[0].cells, values):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _shade_cell(cell, ACCENT_LIGHT)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_run = paragraph.add_run(label + "\n")
        _v1_apply_run_font(label_run)
        label_run.font.size = Pt(9)
        label_run.font.color.rgb = RGBColor.from_string(TEXT_MUTED)
        value_run = paragraph.add_run(value)
        _v1_apply_run_font(value_run)
        value_run.bold = True
        value_run.font.size = Pt(18)
        value_run.font.color.rgb = RGBColor.from_string(_risk_word_color(risk_label) if label == "风险状态" else ACCENT_DARK)
        if note:
            note_run = paragraph.add_run("\n" + note)
            _v1_apply_run_font(note_run)
            note_run.font.size = Pt(8)
            note_run.font.color.rgb = RGBColor.from_string(TEXT_MUTED)
        for margin in ["top", "bottom", "start", "end"]:
            _v1_set_cell_margin(cell, margin, 170)


def _customer_cluster_info_table(
    document: Document,
    context: dict[str, Any],
    cluster: dict[str, Any],
    clusters: list[dict[str, Any]],
    vm_counts: dict[tuple[str, str], int],
) -> None:
    total_current = sum(float((item.get("forecast") or {}).get("current") or 0) for item in clusters)
    total_capacity = sum(float(item.get("total") or 0) for item in clusters)
    risk_clusters = [item for item in clusters if _risk_level(item)[0] != "正常"]
    vm_total = sum(_cluster_vm_count(item, vm_counts.get(_cluster_key(item.get("labels") or {}))) for item in clusters)
    risk = _overall_risk_status(clusters)[0] if clusters else "数据不足"
    tower_count = _tower_count(clusters)
    rows = [
        ("Tower范围", _tower_scope_label(clusters, context["scope_label"])),
        ("Tower数量", f"{tower_count} 个"),
        ("集群范围", _cluster_scope_label(clusters, context["scope_label"])),
        ("集群数量", f"{len(clusters)} 个"),
        ("当前已用容量", _bytes_label(total_current)),
        ("总容量", _bytes_label(total_capacity)),
        ("风险状态", risk),
        ("风险集群数", f"{len(risk_clusters)} 个"),
        ("增长 VM 样本数", f"{vm_total} 台"),
    ]
    table = document.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _v1_set_table_width(table, [3900, 4300])
    _v1_set_table_borders(table, "000000")
    for table_row, (label, value) in zip(table.rows, rows):
        _v1_set_cell(table_row.cells[0], label, fill=ACCENT_LIGHT, color=ACCENT_DARK, bold=True, font_size=10)
        _v1_set_cell(table_row.cells[1], value, color=_risk_word_color(str(value)) if label == "风险状态" else TEXT_DARK, bold=label == "风险状态", font_size=10)


def _customer_cluster_growth_table(document: Document, context: dict[str, Any], clusters: list[dict[str, Any]]) -> None:
    headers = ["Tower", "集群", "近 14 天样本增长", "近 30 天样本增长", "近 90 天样本增长", "90 天预测容量"]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _v1_set_table_width(table, [1400, 1700, 1500, 1500, 1500, 1500])
    _customer_set_blue_headers(table.rows[0].cells, headers)
    if not clusters:
        _v1_set_docx_row(table.add_row().cells, ["当前范围暂无集群容量数据"] + ["-"] * (len(headers) - 1))
        return
    for index, cluster in enumerate(clusters[:8], start=1):
        row = table.add_row().cells
        forecast = cluster.get("forecast") or {}
        labels = cluster.get("labels") or {}
        values = [
            labels.get("tower") or labels.get("tower_id") or "-",
            _cluster_name(cluster),
            _cluster_growth_window_label(cluster, 14),
            _cluster_growth_window_label(cluster, 30),
            _cluster_growth_window_label(cluster, 90),
            _bytes_label(forecast.get("forecast_90d")),
        ]
        _v1_set_docx_row(row, values)
        if index % 2 == 0:
            _v1_shade_row(row, ACCENT_SOFT)
        for cell in row[2:5]:
            _v1_set_cell_text_style(cell, GROWTH_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
    _customer_table_note(document, "说明：近 14/30/90 天样本增长按对应周期内采集样本计算；采集历史不足对应天数时显示数据不足。")


def _cluster_growth_window_label(cluster: dict[str, Any], days: int) -> str:
    growth = _cluster_period_growth_with_min_span(cluster, days)
    return "数据不足" if growth is None else _signed_bytes_label(growth)


def _cluster_period_growth_with_min_span(cluster: dict[str, Any], days: int) -> float | None:
    points = _cluster_points(cluster)
    if len(points) < 2:
        return None
    sample_span_seconds = points[-1][0] - points[0][0]
    if sample_span_seconds < days * 86_400:
        return None
    return _cluster_period_growth(cluster, days)


def _customer_window_days(report: dict[str, Any]) -> int:
    try:
        value = int(report.get("window_days") or (report.get("period_window") or {}).get("days") or 30)
    except (TypeError, ValueError):
        value = 30
    return max(1, value)


def _profile_sample_notice(report: dict[str, Any], profile: ReportPeriodProfile) -> str:
    effective = _effective_report_window(report)
    start = _parse_report_datetime(str(effective.get("start_at") or ""))
    end = _parse_report_datetime(str(effective.get("end_at") or ""))
    if start and end:
        actual_days = max(1, (end.date() - start.date()).days + 1)
        if actual_days < profile.days:
            return f"说明：当前实际采集样本约 {actual_days} 天，{profile.window_growth_label} 已按当前可用样本窗口计算。"
    return f"说明：{profile.window_growth_label} 按当前统计窗口样本计算；如采集历史不足对应天数，则按可用样本窗口计算。"


def _tower_count(clusters: list[dict[str, Any]]) -> int:
    tower_ids: set[str] = set()
    tower_names: set[str] = set()
    for cluster in clusters:
        labels = cluster.get("labels") or {}
        tower_id = str(labels.get("tower_id") or "")
        tower_name = str(labels.get("tower") or "")
        if tower_id:
            tower_ids.add(tower_id)
        elif tower_name:
            tower_names.add(tower_name)
    return len(tower_ids or tower_names)


def _customer_usage_bars(document: Document, clusters: list[dict[str, Any]]) -> None:
    if not clusters:
        _customer_table_note(document, "当前范围暂无集群容量数据。")
        return
    cluster = max(clusters, key=_cluster_used_ratio)
    total = _float_or_none(cluster.get("total"))
    if total is None or total <= 0:
        _customer_table_note(document, f"{_cluster_full_name(cluster)} 暂无总容量数据，无法生成容量使用率可视化。")
        return
    current_ratio = _cluster_used_ratio(cluster)
    future = _float_or_none((cluster.get("forecast") or {}).get("forecast_90d"))
    current = _float_or_none((cluster.get("forecast") or {}).get("current"))
    future_ratio = (future if future is not None else current or 0) / total
    threshold = _float_or_none(cluster.get("warning")) or total
    _customer_table_note(document, f"展示对象：{_cluster_full_name(cluster)}。多集群范围下默认展示当前使用率最高的集群。")
    _customer_usage_bar(document, "当前使用率", current_ratio, ACCENT_DARK, threshold)
    _customer_usage_bar(document, "90 天预测使用率", future_ratio, GROWTH_BLUE, threshold)
    _customer_table_note(document, _customer_usage_summary(cluster, current_ratio, future_ratio))


def _customer_usage_bar(document: Document, label: str, ratio: float, color: str, threshold: float) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(f"{label}   {_percent_label(ratio)}")
    _v1_apply_run_font(run)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor.from_string(TEXT_DARK)

    bar = document.add_paragraph()
    bar.paragraph_format.space_before = Pt(0)
    bar.paragraph_format.space_after = Pt(4)
    bar.paragraph_format.keep_with_next = True
    normalized = max(0.0, min(ratio, 1.0))
    width = 50
    used = max(0, min(width, round(normalized * 100)))
    blocks = "█" * used + "░" * (width - used)
    bar_run = bar.add_run(f"  {blocks}")
    _v1_apply_run_font(bar_run)
    bar_run.font.size = Pt(8)
    bar_run.font.color.rgb = RGBColor.from_string(color)
    threshold_run = bar.add_run(f"   | 容量阈值   {_bytes_label(threshold)}")
    _v1_apply_run_font(threshold_run)
    threshold_run.font.size = Pt(9)
    threshold_run.font.color.rgb = RGBColor.from_string(TEXT_MUTED)


def _customer_usage_summary(cluster: dict[str, Any], current_ratio: float, future_ratio: float) -> str:
    cluster_name = _cluster_full_name(cluster)
    if current_ratio >= 0.8:
        return f"{cluster_name} 当前使用率已达到高风险阈值，容量安全边际不足，建议优先确认扩容周期和可清理空间。"
    if future_ratio >= 0.8:
        return f"{cluster_name} 当前使用率尚未超过高风险阈值，但 90 天预测接近或超过 80%，容量安全边际收窄，建议进入容量跟踪。"
    if future_ratio >= 0.75:
        return f"{cluster_name} 90 天预测使用率接近关注阈值，容量安全边际需持续观察。"
    return f"{cluster_name} 使用率较低，预测趋势线与当前基线接近，容量安全边际较充裕。"


def _customer_vm_window_note(document: Document, context: dict[str, Any], vms: list[dict[str, Any]], suffix: str) -> None:
    paragraph = document.add_paragraph(f"统计窗口：{_vm_sample_window_label(vms, context['report'])} | {suffix}")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.keep_with_next = True
    for run in paragraph.runs:
        _v1_apply_run_font(run)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(TEXT_MUTED)


def _customer_vm_table(document: Document, vms: list[dict[str, Any]], sort_mode: str, *, empty_text: str) -> None:
    headers = ["排名", "虚拟机名称", "当前容量", "期初容量", "增长量", "增长率"]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _v1_set_table_width(table, [900, 2200, 1500, 1500, 1500, 1300])
    _customer_set_blue_headers(table.rows[0].cells, headers)
    _repeat_table_header(table.rows[0])
    if not vms:
        row = table.add_row().cells
        _v1_set_docx_row(row, [empty_text] + ["-"] * (len(headers) - 1))
        _prevent_row_split(table.rows[-1])
        return
    for index, vm in enumerate(vms, start=1):
        row = table.add_row().cells
        labels = vm.get("labels") or {}
        values = [
            index,
            labels.get("vm") or labels.get("vm_name") or labels.get("vm_id") or "-",
            _bytes_label((vm.get("forecast") or {}).get("current")),
            _bytes_label(vm.get("previous_value")),
            _bytes_label(vm.get("growth_amount")),
            _percent_label(vm.get("growth_ratio")),
        ]
        _v1_set_docx_row(row, values)
        if index % 2 == 0:
            _v1_shade_row(row, ACCENT_SOFT)
        _v1_set_cell_text_style(row[0], TEXT_MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
        _v1_set_cell_text_style(row[4], GROWTH_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
        ratio = float(vm.get("growth_ratio") or 0)
        ratio_color = RATIO_RED if ratio >= 0.5 else RATIO_ORANGE if ratio >= 0.2 else TEXT_DARK
        _v1_set_cell_text_style(row[5], ratio_color, align=WD_ALIGN_PARAGRAPH.CENTER)
        _prevent_row_split(table.rows[-1])


def _customer_set_blue_headers(cells: Any, headers: list[str]) -> None:
    for cell, header in zip(cells, headers):
        _v1_set_cell(cell, header, fill=ACCENT, color="FFFFFF", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font_size=9)


def _customer_table_note(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(8)
    for run in paragraph.runs:
        _v1_apply_run_font(run)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(TEXT_MUTED)


def _customer_add_emphasis_text(paragraph: Any, text: str, *, base_size: float = 10.5) -> None:
    for segment, highlight in _customer_emphasis_segments(text):
        run = paragraph.add_run(_customer_emphasis_text(segment, highlight))
        _v1_apply_run_font(run)
        run.font.size = Pt(base_size + 1 if highlight else base_size)
        run.font.color.rgb = RGBColor.from_string(TEXT_DARK)
        run.bold = bool(highlight)


def _customer_emphasis_text(segment: str, highlight: bool) -> str:
    if highlight and _customer_is_vm_name_segment(segment):
        return f" {segment} "
    return segment


def _customer_is_vm_name_segment(segment: str) -> bool:
    if _customer_is_numeric_segment(segment) or _customer_is_scope_segment(segment):
        return False
    vm_patterns = [
        r"[^，。！？（）]+?(?=\s*(?:增长|当前容量|在统计窗口|（))",
        r"Window\s+VM\s+[A-Za-z0-9_-]+",
        r"Day\s+VM",
        r"VM\s+[A-Za-z0-9_-]+",
    ]
    return any(re.fullmatch(pattern, segment) for pattern in vm_patterns) or bool(segment.strip())


def _customer_is_numeric_segment(segment: str) -> bool:
    numeric_patterns = [
        r"[+-]?\d+(?:\.\d+)?\s*(?:B|KB|MB|GB|TB|PB)",
        r"\d+(?:\.\d+)?%",
        r"\d+\s*天",
    ]
    return any(re.fullmatch(pattern, segment) for pattern in numeric_patterns)


def _customer_is_scope_segment(segment: str) -> bool:
    scope_patterns = [
        r"全部\s*Tower（\d+\s*个）",
        r"全部集群（\d+\s*个）",
        r"Tower\s+[A-Za-z0-9_-]+",
        r"Cluster\s+[A-Za-z0-9_-]+",
        r"[A-Za-z0-9_-]+(?:\s*/\s*[A-Za-z0-9_-]+)+",
    ]
    return any(re.fullmatch(pattern, segment) for pattern in scope_patterns)


def _customer_emphasis_segments(text: str) -> list[tuple[str, bool]]:
    patterns = [
        r"[+-]\d+(?:\.\d+)?\s*(?:B|KB|MB|GB|TB|PB)",
        r"\d+(?:\.\d+)?\s*(?:B|KB|MB|GB|TB|PB)",
        r"\d+(?:\.\d+)?%",
        r"\d+\s*天",
        r"全部\s*Tower（\d+\s*个）",
        r"全部集群（\d+\s*个）",
        r"[A-Z0-9][A-Za-z0-9_-]+(?:\s*/\s*[A-Z0-9][A-Za-z0-9_-]+)+",
        r"Tower\s+[A-Za-z0-9_-]+",
        r"Cluster\s+[A-Za-z0-9_-]+",
        r"(?<=下的 )[^，。！？（）]+?(?=\s*(?:增长|当前容量|在统计窗口|（))",
        r"Window\s+VM\s+[A-Za-z0-9_-]+",
        r"Day\s+VM",
        r"VM\s+[A-Za-z0-9_-]+",
    ]
    combined = re.compile("|".join(f"({pattern})" for pattern in patterns))
    segments: list[tuple[str, bool]] = []
    cursor = 0
    for match in combined.finditer(text):
        if match.start() > cursor:
            segments.append((text[cursor:match.start()], False))
        segments.append((match.group(0), True))
        cursor = match.end()
    if cursor < len(text):
        segments.append((text[cursor:], False))
    return segments or [(text, False)]


def _customer_add_figure(document: Document, image: BytesIO, caption: str, width: float) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(caption)
    _v1_apply_run_font(run)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(ACCENT_DARK)
    picture = document.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.add_run().add_picture(image, width=Inches(width))


def _customer_risk_matrix_table(document: Document, clusters: list[dict[str, Any]], top_vms: list[dict[str, Any]]) -> None:
    headers = ["风险项", "当前状态", "风险等级", "评估说明"]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _v1_set_table_width(table, [1600, 1500, 1200, 4400])
    _customer_set_blue_headers(table.rows[0].cells, headers)
    for index, row_data in enumerate(_customer_risk_matrix_rows(clusters, top_vms), start=1):
        row = table.add_row().cells
        _v1_set_docx_row(row, [row_data["item"], row_data["status"], row_data["level"], row_data["description"]])
        if index % 2 == 0:
            _v1_shade_row(row, ACCENT_SOFT)
        _v1_set_cell_text_style(row[2], _risk_text_color(row_data["level"]), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)


def _v1_setup_document(document: Document, context: dict[str, Any]) -> None:
    section = document.sections[0]
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    section.top_margin = Inches(0.66)
    section.bottom_margin = Inches(0.64)
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = document.styles[style_name]
        _v1_apply_style_font(style)
    document.styles["Normal"].font.size = Pt(9)
    document.styles["Heading 1"].font.size = Pt(15)
    document.styles["Heading 1"].font.bold = True
    document.styles["Heading 1"].font.color.rgb = RGBColor.from_string(ACCENT_DARK)
    document.styles["Heading 2"].font.size = Pt(11)
    document.styles["Heading 2"].font.bold = True
    document.styles["Heading 2"].font.color.rgb = RGBColor.from_string(ACCENT_DARK)
    _v1_setup_header_footer(section, context)


def _v1_apply_style_font(style: Any) -> None:
    style.font.name = DOCX_FONT_ASCII
    style._element.rPr.rFonts.set(qn("w:ascii"), DOCX_FONT_ASCII)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), DOCX_FONT_ASCII)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), DOCX_FONT_EAST_ASIA)


def _v1_apply_run_font(run: Any) -> None:
    run.font.name = DOCX_FONT_ASCII
    run._element.rPr.rFonts.set(qn("w:ascii"), DOCX_FONT_ASCII)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), DOCX_FONT_ASCII)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), DOCX_FONT_EAST_ASIA)


def _v1_setup_header_footer(section: Any, context: dict[str, Any], footer_label: str | None = None) -> None:
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("SmartX Storage Forecast | 容量预测报告")
    _v1_apply_run_font(run)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(TEXT_MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f"{footer_label or context['scope_label']} · {context['generated_at']}")
    _v1_apply_run_font(run)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(TEXT_MUTED)


def _v1_start_section(document: Document, context: dict[str, Any], footer_label: str) -> None:
    section = document.add_section(WD_SECTION.NEW_PAGE)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    section.top_margin = Inches(0.66)
    section.bottom_margin = Inches(0.64)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    _v1_clear_paragraph(section.header.paragraphs[0])
    _v1_clear_paragraph(section.footer.paragraphs[0])
    _v1_setup_header_footer(section, context, footer_label=footer_label)


def _v1_clear_paragraph(paragraph: Any) -> None:
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)


def _v1_add_cover_brand(document: Document) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _v1_set_table_width(table, [900, 7200])
    _v1_set_table_borders(table, "FFFFFF")
    _v1_set_cell(table.rows[0].cells[0], "", fill=ACCENT)
    cell = table.rows[0].cells[1]
    _v1_set_cell(cell, REPORT_PRODUCT_NAME, fill=ACCENT_SOFT, color=ACCENT_DARK, bold=True, font_size=9)
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _v1_add_cover_rule(document: Document) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _v1_set_table_width(table, [7800])
    _v1_set_table_borders(table, "FFFFFF")
    _shade_cell(table.rows[0].cells[0], ACCENT)


def _v1_add_callout(document: Document, title: str, body: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    _v1_set_table_width(table, [8200])
    _v1_set_table_borders(table, BORDER)
    cell = table.rows[0].cells[0]
    _shade_cell(cell, ACCENT_SOFT)
    paragraph = cell.paragraphs[0]
    title_run = paragraph.add_run(f"{title}：")
    _v1_apply_run_font(title_run)
    title_run.bold = True
    title_run.font.size = Pt(9)
    title_run.font.color.rgb = RGBColor.from_string(ACCENT_DARK)
    body_run = paragraph.add_run(body)
    _v1_apply_run_font(body_run)
    body_run.font.size = Pt(9)
    body_run.font.color.rgb = RGBColor.from_string(TEXT_DARK)
    document.add_paragraph()


def _v1_add_figure(document: Document, image: BytesIO, caption: str, width: float) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run(caption)
    _v1_apply_run_font(run)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(TEXT_DARK)
    paragraph.paragraph_format.space_after = Pt(2)
    picture_paragraph = document.add_paragraph()
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.paragraph_format.keep_together = True
    picture_paragraph.add_run().add_picture(image, width=Inches(width))
    document.add_paragraph()


def _v1_add_cover(document: Document, context: dict[str, Any]) -> None:
    _v1_add_cover_brand(document)
    document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(REPORT_COVER_TITLE)
    _v1_apply_run_font(run)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor.from_string(ACCENT_DARK)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(REPORT_COVER_SUBTITLE)
    _v1_apply_run_font(run)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(TEXT_MUTED)

    _v1_add_cover_rule(document)
    document.add_paragraph()

    table = document.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    _v1_set_table_width(table, [2300, 5400])
    _v1_set_table_borders(table, BORDER)
    rows = [
        ("导出范围", context["scope_label"]),
        ("生成时间", context["generated_at"]),
        ("统计窗口", _v1_vm_window_label(context)),
        ("预测窗口", f"预测 {context['report'].get('forecast_days') or 90} 天"),
        ("集群数量", f"{len(context['clusters'])} 个"),
        ("当前软件版本", context.get("app_version", "-")),
    ]
    for row, (key, value) in zip(table.rows, rows):
        _v1_set_cell(row.cells[0], key, fill=ACCENT, color="FFFFFF", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _v1_set_cell(row.cells[1], value, fill=ACCENT_SOFT, color=TEXT_DARK)

    note = document.add_paragraph(f"本报告由 {REPORT_PRODUCT_NAME} {context.get('app_version', '-')} 自动生成")
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in note.runs:
        _v1_apply_run_font(run)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(TEXT_MUTED)

    document.add_paragraph()
    _v1_add_callout(
        document,
        "报告说明",
        "本报告基于平台采集容量样本自动生成，用于客户容量趋势复盘、风险沟通与扩容规划参考。预测结果应结合业务上线计划和实际资源治理策略共同判断。",
    )
    document.add_page_break()


def _v1_add_report_overview(document: Document, context: dict[str, Any]) -> None:
    clusters = context["clusters"]
    report = context["report"]
    top_vms = _customer_growth_vms(report)
    total_current = sum(float((cluster.get("forecast") or {}).get("current") or 0) for cluster in clusters)
    total_month = sum(_cluster_period_growth(cluster, 30) for cluster in clusters)
    total_forecast_90 = sum(float((cluster.get("forecast") or {}).get("forecast_90d") or 0) for cluster in clusters)
    total_capacity = sum(float(cluster.get("total") or 0) for cluster in clusters)
    total_vm_count = _report_vm_count(report)
    risk_label, risk_fill, risk_note = _overall_risk_status(clusters)
    document.add_heading("一  执行摘要", level=1)
    paragraph = document.add_paragraph("报告摘要")
    for run in paragraph.runs:
        _v1_apply_run_font(run)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(ACCENT_DARK)
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _v1_set_table_width(table, [2050, 2050, 2050, 2050])
    _v1_set_table_borders(table, BORDER)
    values = [
        ("当前已用容量", _bytes_label(total_current), f"总容量 {_bytes_label(total_capacity)}"),
        ("较上月增长", _signed_bytes_label(total_month), ""),
        ("90天预测容量", _bytes_label(total_forecast_90), f"预计使用率 {_percent_label(total_forecast_90 / total_capacity) if total_capacity > 0 else '-'}"),
        ("风险状态", risk_label, f"虚拟机数量 {total_vm_count} 台"),
    ]
    for cell, (label, value, note) in zip(table.rows[0].cells, values):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _shade_cell(cell, risk_fill if label == "风险状态" else ACCENT_LIGHT)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(label + "\n")
        _v1_apply_run_font(run)
        run.bold = False
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(TEXT_MUTED)
        value_run = paragraph.add_run(value)
        _v1_apply_run_font(value_run)
        value_run.bold = True
        value_run.font.size = Pt(13)
        value_run.font.color.rgb = RGBColor.from_string(_risk_word_color(risk_label) if label == "风险状态" else ACCENT_DARK)
        if note:
            note_run = paragraph.add_run("\n" + note)
            _v1_apply_run_font(note_run)
            note_run.font.size = Pt(8)
            note_run.font.color.rgb = RGBColor.from_string(TEXT_MUTED)
    risk_count = sum(1 for cluster in clusters if _risk_level(cluster)[0] != "正常")
    paragraph = document.add_paragraph(f"容量关注集群：{risk_count} 个")
    for run in paragraph.runs:
        _v1_apply_run_font(run)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(TEXT_MUTED)
    _v1_add_callout(document, "容量风险摘要", _capacity_risk_summary(context["report"]))
    _v1_add_callout(document, "结论摘要", _overview_sentence(clusters))
    _v1_add_key_findings(document, clusters, top_vms, risk_note, context)


def _v1_add_cluster_directory(document: Document, clusters: list[dict[str, Any]]) -> None:
    document.add_heading("目录", level=1)
    entries: list[tuple[str, str, str, str | None]] = [
        ("1", "执行摘要（报告摘要）", "1", None),
        ("2", "集群容量增长概览", "1.1", None),
        ("3", "本次统计窗口增长 VM", "1.2", None),
        ("4", "风险评估与建议", "1.3", None),
    ]
    for index, cluster in enumerate(clusters, start=1):
        labels = cluster.get("labels", {})
        cluster_title = _v1_cluster_title(labels)
        entries.extend(
            [
                (str(len(entries) + 1), cluster_title, f"2.{index}", _v1_cluster_bookmark(index)),
                (str(len(entries) + 1), f"{cluster_title} - Top 10 VM 增长量", f"2.{index}.1", _v1_cluster_bookmark(index)),
                (str(len(entries) + 1), f"{cluster_title} - 增长量 TOP100 虚拟机", f"2.{index}.2", _v1_cluster_bookmark(index)),
                (str(len(entries) + 1), f"{cluster_title} - 增长率 TOP100 虚拟机", f"2.{index}.3", _v1_cluster_bookmark(index)),
            ]
        )
    if not clusters:
        entries.append(("4", "当前导出范围内暂无集群", "-", None))

    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _v1_set_table_width(table, [900, 4300, 2500])
    _v1_set_docx_headers(table.rows[0].cells, ["序号", "内容", "章节"])
    _v1_set_table_borders(table, BORDER)
    for row_index, (order, title, chapter, bookmark) in enumerate(entries, start=1):
        row = table.add_row().cells
        _v1_set_docx_row(row, [order, "", chapter])
        if bookmark:
            _v1_add_internal_link(row[1].paragraphs[0], bookmark, title)
        else:
            _v1_set_cell(row[1], title)
        if row_index % 2 == 0:
            _v1_shade_row(row, ACCENT_SOFT)
    _v1_add_callout(document, "定位说明", "可通过本目录快速确认集群章节编号；Word 打开后也可使用导航窗格按标题跳转。")
    document.add_page_break()


def _v1_add_key_findings(document: Document, clusters: list[dict[str, Any]], top_vms: list[dict[str, Any]], risk_note: str, context: dict[str, Any]) -> None:
    document.add_heading("关键发现", level=2)
    findings = _customer_key_findings(clusters, top_vms, risk_note, context)
    for finding in findings:
        paragraph = document.add_paragraph(finding)
        paragraph.paragraph_format.left_indent = Inches(0.16)
        for run in paragraph.runs:
            _v1_apply_run_font(run)
            run.font.size = Pt(9)
    document.add_paragraph()


def _v1_add_risk_and_advice(document: Document, report: dict[str, Any]) -> None:
    clusters = report.get("clusters") or []
    top_vms = _customer_growth_vms(report)
    document.add_heading("四  风险评估与建议", level=1)
    document.add_heading("4.1  容量风险评估矩阵", level=2)
    _v1_add_risk_matrix_table(document, clusters, top_vms)
    document.add_heading("4.2  运维建议", level=2)
    for title, items in _customer_operation_advice(clusters, top_vms):
        _v1_add_advice_group(document, title, items)
    _v1_add_callout(
        document,
        "声明",
        "本报告基于平台采集的容量数据自动生成，预测结果基于历史增长趋势推算，仅供参考。实际容量变化受业务负载、数据写入模式等因素影响。",
    )


def _v1_add_risk_matrix_table(document: Document, clusters: list[dict[str, Any]], top_vms: list[dict[str, Any]]) -> None:
    rows = _customer_risk_matrix_rows(clusters, top_vms)
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _v1_set_table_width(table, [1700, 1600, 1300, 3600])
    _v1_set_docx_headers(table.rows[0].cells, ["风险项", "当前状态", "风险等级", "评估说明"])
    _v1_set_table_borders(table, BORDER)
    for index, row_data in enumerate(rows, start=1):
        row = table.add_row().cells
        _v1_set_docx_row(row, [row_data["item"], row_data["status"], row_data["level"], row_data["description"]])
        if index % 2 == 0:
            _v1_shade_row(row, ACCENT_SOFT)
        _set_cell_text_color(row[2], _risk_text_color(row_data["level"]))


def _v1_add_advice_group(document: Document, title: str, items: list[str]) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(title)
    _v1_apply_run_font(run)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(GROWTH_BLUE)
    for item in items:
        body = document.add_paragraph(item)
        body.paragraph_format.left_indent = Inches(0.28)
        for run in body.runs:
            _v1_apply_run_font(run)
            run.font.size = Pt(9)


def _v1_add_cluster_growth_chart(document: Document, clusters: list[dict[str, Any]]) -> None:
    trend_chart = _scope_trend_line_chart(clusters, "容量使用率趋势")
    if trend_chart is not None:
        _v1_add_figure(document, trend_chart, "图 1：容量使用趋势", width=6.6)
    else:
        document.add_paragraph("暂无足够历史数据生成容量趋势图。")
    top_chart = _cluster_top_growth_bar_chart(clusters, "Top 5 集群月增长量")
    if top_chart is not None:
        _v1_add_figure(document, top_chart, "图 2：Top 5 集群月增长量", width=6.8)


def _v1_add_cluster_summary_table(document: Document, clusters: list[dict[str, Any]], vm_counts: dict[tuple[str, str], int] | None = None) -> None:
    headers = ["Tower", "集群", "虚拟机数量", "当前容量", "较上月", "较上季度", "较上一年", "90天预测容量", "风险"]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _v1_set_table_width(table, [1000, 1550, 900, 1100, 1050, 1050, 1050, 1100, 800])
    _v1_set_docx_headers(table.rows[0].cells, headers)
    _v1_set_table_borders(table, BORDER)
    if not clusters:
        _v1_set_docx_row(table.add_row().cells, ["当前范围暂无集群容量数据"] + ["-"] * (len(headers) - 1))
        return
    for index, cluster in enumerate(clusters, start=1):
        labels = cluster.get("labels", {})
        forecast = cluster.get("forecast", {})
        risk, fill = _risk_level(cluster)
        vm_count = _cluster_vm_count(cluster, (vm_counts or {}).get(_cluster_key(labels)))
        row = table.add_row().cells
        values = [
            labels.get("tower") or labels.get("tower_id") or "-",
            labels.get("cluster") or labels.get("cluster_id") or "-",
            f"{vm_count} 台",
            _bytes_label(forecast.get("current")),
            _signed_bytes_label(_cluster_period_growth(cluster, 30)),
            _signed_bytes_label(_cluster_period_growth(cluster, 90)),
            _signed_bytes_label(_cluster_period_growth(cluster, 365)),
            _bytes_label(forecast.get("forecast_90d")),
            risk,
        ]
        _v1_set_docx_row(row, values)
        if index % 2 == 0:
            _v1_shade_row(row, ACCENT_SOFT)
        for amount_cell in row[4:7]:
            _v1_set_cell_text_style(amount_cell, GROWTH_BLUE, bold=False)
        _v1_set_cell_text_style(row[-1], _risk_word_color(risk), bold=True)
        _shade_cell(row[-1], fill)


def _v1_add_cluster_customer_summary(document: Document, cluster: dict[str, Any], vm_count: int, *, visible_vm_count: int | None = None) -> None:
    forecast = cluster.get("forecast", {})
    risk, fill = _risk_level(cluster)
    total_vm_count = _cluster_vm_count(cluster, visible_vm_count if visible_vm_count is not None else vm_count)
    table = document.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    _v1_set_table_width(table, [2100, 2100, 2100, 2100])
    items = [
        ("当前已用", _bytes_label(forecast.get("current"))),
        ("较上月增加", _signed_bytes_label(_cluster_period_growth(cluster, 30))),
        ("较上季度增加", _signed_bytes_label(_cluster_period_growth(cluster, 90))),
        ("较上一年增加", _signed_bytes_label(_cluster_period_growth(cluster, 365))),
        ("90天预测容量", _bytes_label(forecast.get("forecast_90d"))),
        ("虚拟机数量", f"{total_vm_count} 台"),
        ("风险状态", risk),
        ("增长VM样本", f"{vm_count} 台"),
    ]
    for cell, (label, value) in zip([cell for row in table.rows for cell in row.cells], items):
        _shade_cell(cell, fill if label == "风险状态" else ACCENT_LIGHT)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(label + "\n")
        _v1_apply_run_font(run)
        run.bold = False
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(TEXT_MUTED)
        run = paragraph.add_run(value)
        _v1_apply_run_font(run)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor.from_string(_risk_word_color(risk) if label == "风险状态" else ACCENT_DARK)
    document.add_paragraph(_v1_cluster_advice(cluster))


def _v1_cluster_advice(cluster: dict[str, Any]) -> str:
    risk, _ = _risk_level(cluster)
    month = _cluster_period_growth(cluster, 30)
    quarter = _cluster_period_growth(cluster, 90)
    if risk == "高风险":
        return "建议：当前容量已达到关注阈值，建议尽快确认可用容量、扩容周期和重点增长虚拟机。"
    if risk == "需关注":
        return "建议：90 天预测接近容量阈值，建议跟踪增长趋势并提前准备扩容或清理方案。"
    if quarter > month * 2.5 and quarter > 0:
        return "建议：季度增长明显高于月增长节奏，建议复核近期业务上线或批量数据写入情况。"
    return "建议：当前容量趋势相对平稳，建议持续观察 Top 增长虚拟机并保持月度复盘。"


def _v1_add_single_cluster_charts(document: Document, cluster: dict[str, Any], vms: list[dict[str, Any]], vm_chart_title: str = "Top 10 VM 增长量") -> None:
    trend_chart = _cluster_trend_line_chart(cluster, "集群容量使用趋势")
    if trend_chart is not None:
        _v1_add_figure(document, trend_chart, "容量使用趋势", width=5.9)
    top_chart = _vm_top_growth_bar_chart(vms, vm_chart_title)
    if top_chart is not None:
        _v1_add_figure(document, top_chart, vm_chart_title, width=6.4)


def _v1_add_vm_table(document: Document, vms: list[dict[str, Any]], sort_mode: str, include_cluster: bool = False, empty_text: str = "暂无增长数据") -> None:
    headers = _v1_vm_headers(include_cluster=include_cluster, sort_mode=sort_mode)
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    widths = [560, 1050, 1350, 2600, 1050, 1050, 1050, 850] if include_cluster else [576, 3312, 1224, 1224, 1224, 1080]
    _v1_set_table_width(table, widths)
    _v1_set_docx_headers(table.rows[0].cells, headers)
    _v1_set_table_borders(table, BORDER)
    if not vms:
        empty_row = table.add_row().cells
        _v1_set_docx_row(empty_row, [empty_text] + ["-"] * (len(headers) - 1))
        _v1_shade_row(empty_row, ACCENT_SOFT)
        return
    for index, vm in enumerate(vms, start=1):
        row = table.add_row().cells
        _v1_set_docx_row(row, _v1_vm_row(vm, include_cluster=include_cluster, raw=False, rank=index))
        if index % 2 == 0:
            _v1_shade_row(row, ACCENT_SOFT)
        _v1_style_vm_docx_row(row, vm, include_cluster=include_cluster)


def _v1_add_vm_window_note(document: Document, context: dict[str, Any], vms: list[dict[str, Any]] | None = None) -> None:
    note = document.add_paragraph(f"统计窗口：{_vm_sample_window_label(vms or [], context['report'])}")
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in note.runs:
        _v1_apply_run_font(run)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(TEXT_MUTED)


def _v1_set_docx_headers(cells: Any, headers: list[str]) -> None:
    for cell, header in zip(cells, headers):
        _v1_set_cell(cell, header, fill=ACCENT, color="FFFFFF", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)


def _v1_set_docx_row(cells: Any, values: list[Any]) -> None:
    for cell, value in zip(cells, values):
        _v1_set_cell(cell, value)


def _v1_style_vm_docx_row(cells: Any, vm: dict[str, Any], *, include_cluster: bool) -> None:
    amount_index = 6 if include_cluster else 4
    ratio_index = 7 if include_cluster else 5
    vm_index = 3 if include_cluster else 1
    numeric_start = 4 if include_cluster else 2

    _v1_set_cell_text_style(cells[0], TEXT_MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
    _v1_set_cell_text_style(cells[vm_index], TEXT_DARK, align=WD_ALIGN_PARAGRAPH.LEFT)
    for cell in cells[numeric_start:]:
        _v1_set_cell_text_style(cell, TEXT_DARK, align=WD_ALIGN_PARAGRAPH.CENTER)
    _v1_set_cell_text_style(cells[amount_index], GROWTH_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)

    ratio = float(vm.get("growth_ratio") or 0)
    ratio_color = RATIO_RED if ratio >= 0.5 else RATIO_ORANGE if ratio >= 0.2 else TEXT_DARK
    _v1_set_cell_text_style(cells[ratio_index], ratio_color, align=WD_ALIGN_PARAGRAPH.CENTER)


def _v1_set_cell_text_style(cell: Any, color: str, *, bold: bool | None = None, align: int | None = None) -> None:
    for paragraph in cell.paragraphs:
        if align is not None:
            paragraph.alignment = align
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor.from_string(color)
            if bold is not None:
                run.bold = bold


def _v1_set_table_borders(table: Any, color: str) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _repeat_table_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def _prevent_row_split(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def _v1_shade_row(cells: Any, fill: str) -> None:
    for cell in cells:
        _shade_cell(cell, fill)


def _v1_set_cell(
    cell: Any,
    value: Any,
    fill: str | None = None,
    color: str = TEXT_DARK,
    bold: bool = False,
    align: int | None = None,
    font_size: int = 8,
) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        _shade_cell(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(str(value))
    _v1_apply_run_font(run)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor.from_string(color)
    for margin in ["top", "start", "bottom", "end"]:
        _v1_set_cell_margin(cell, margin, 110)


def _v1_add_bookmarked_heading(document: Document, text: str, bookmark_name: str, level: int, bookmark_id: int) -> None:
    paragraph = document.add_heading(level=level)
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), bookmark_name)
    paragraph._p.append(start)
    run = paragraph.add_run(text)
    _v1_apply_run_font(run)
    if level == 1:
        run.bold = True
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor.from_string(ACCENT_DARK)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.append(end)


def _v1_add_internal_link(paragraph: Any, bookmark_name: str, text: str) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)
    run_element = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), ACCENT)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(color)
    properties.append(underline)
    run_element.append(properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run_element.append(text_element)
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)
    for run in paragraph.runs:
        _v1_apply_run_font(run)


def _v1_set_cell_margin(cell: Any, margin: str, size: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    node = tc_mar.find(qn(f"w:{margin}"))
    if node is None:
        node = OxmlElement(f"w:{margin}")
        tc_mar.append(node)
    node.set(qn("w:w"), str(size))
    node.set(qn("w:type"), "dxa")


def _v1_set_table_width(table: Any, widths: list[int]) -> None:
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Pt(width / 20)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def _v1_vm_headers(include_cluster: bool, sort_mode: str) -> list[str]:
    headers = ["排名"]
    if include_cluster:
        headers.extend(["Tower", "集群"])
    headers.extend(["虚拟机名称", "当前容量", "期初容量", "增长量", "增长率"])
    if sort_mode == "ratio":
        headers[-1] = "增长率 ↓"
    else:
        headers[-2] = "增长量 ↓"
    return headers


def _v1_vm_row(vm: dict[str, Any], include_cluster: bool, raw: bool = False, rank: int | None = None) -> list[Any]:
    labels = vm.get("labels", {})
    forecast = vm.get("forecast", {})
    current = forecast.get("current") or 0
    previous = vm.get("previous_value") or 0
    amount = vm.get("growth_amount") or 0
    ratio = vm.get("growth_ratio") or 0
    row: list[Any] = [rank or ""]
    if include_cluster:
        row.extend([labels.get("tower") or labels.get("tower_id") or "-", labels.get("cluster") or labels.get("cluster_id") or "-"])
    row.extend([
        labels.get("vm") or labels.get("vm_name") or labels.get("vm_id") or "-",
        current if raw else _bytes_label(current),
        previous if raw else _bytes_label(previous),
        amount if raw else _bytes_label(amount),
        ratio if raw else _percent_label(ratio),
    ])
    return row


def _v1_cluster_title(labels: dict[str, Any]) -> str:
    tower = labels.get("tower")
    cluster = labels.get("cluster") or labels.get("cluster_id") or "未知集群"
    return f"{tower} - {cluster}" if tower else str(cluster)


def _v1_cluster_footer_label(labels: dict[str, Any]) -> str:
    tower = labels.get("tower") or labels.get("tower_id") or "Tower"
    cluster = labels.get("cluster") or labels.get("cluster_id") or "集群"
    return f"{tower} - {cluster}"


def _v1_cluster_bookmark(index: int) -> str:
    return f"cluster_{index}"


def _v1_vm_window_label(context: dict[str, Any]) -> str:
    return f"统计窗口：{_period_window_label(context['report'])}"


def _vm_growth_bucket_label(report: dict[str, Any]) -> str:
    bucket = report.get("vm_growth_sample_bucket") or {}
    min_days = bucket.get("min_days")
    max_days = bucket.get("max_days")
    if min_days is not None and max_days is not None:
        return f"VM 增长样本跨度：{min_days} 天 <= 样本跨度 <= {max_days} 天"
    return f"VM 增长样本跨度：当前导出周期 {report.get('window_days') or 30} 天"
